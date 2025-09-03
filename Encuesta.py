############ Instalación de paquetes
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
import matplotlib.pyplot as plt
import numpy as np
import os
import io  # ← aquí
from io import BytesIO
import streamlit as st
import pandas as pd
from datetime import datetime
import uuid  # Para generar IDs únicos
#import openpyxl
from datetime import date
import json
import re
from collections import defaultdict
import yagmail
##############################
from docx import Document
import streamlit as st
import io
import gspread
from oauth2client.service_account import ServiceAccountCredentials
##############################################
def set_cell_background(cell, rgb_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), rgb_color)
    tcPr.append(shd)

# Diccionario para los textos de valoración
texto_valoracion = {
    1: "1. No cumple",
    2: "2. Incipiente",
    3: "3. Aceptable",
    4: "4. Satisfactorio",
    5: "5. Óptimo"
}

texto_valoracion_cond = {
    1: "1. Cumple de forma incipiente uno o dos criterios",
    2: "2. Cumple de forma incipiente uno o dos criterios",
    3: "3. Cumple de forma aceptable mínimo tres criterios",
    4: "4. Cumple de forma satisfactoria mínimo tres criterios",
    5: "5. Cumple de forma óptima todos los criterios"
}
##############################################
def tabla_detalle_condiciones(doc, dimensiones, nombres_subdimensiones, st_session_state):
    # Recorre cada subdimensión (ejemplo: D1.1)
    for subdim, variables in dimensiones.items():
        if subdim not in nombres_subdimensiones:
            continue

        # Crea tabla: filas = 1 encabezado + 4 criterios + 1 total
        table = doc.add_table(rows=1+4+1, cols=4)
        table.style = 'Table Grid'
        # Encabezados
        hdr = ["CONDICIONES DE LOS SERVICIOS DE REHABILITACIÓN", "CRITERIOS", "VALORACIÓN DE CRITERIOS", "VALORACIÓN DE LA CONDICIÓN"]
        for i, h in enumerate(hdr):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            if i == 0:
                set_cell_background(cell, "FFF6D9")  # fondo amarillo claro
            elif i == 3:
                set_cell_background(cell, "000000")  # fondo negro

        # Criterios y valoraciones
        preguntas = [
            "La institución presta servicio de psicología y/o trabajo social.",
            "La institución presta servicios de fisioterapia, fonoaudiología y/o terapia ocupacional.",
            "Los servicios de rehabilitación disponibles corresponden con el nivel de complejidad.",
            "Los servicios de rehabilitación se organizan en un área específica de la institución.",
        ]
        # Para generalizar, usa tus textos reales de preguntas según subdimensión

        for i in range(4):
            row = table.rows[i+1]
            if i == 0:
                # Celda de condición con nombre largo
                row.cells[0].text = nombres_subdimensiones[subdim]
            else:
                row.cells[0].text = ""
            row.cells[1].text = preguntas[i]
            # Valoración del criterio
            val_key = f"p{subdim.replace('.', '_')}_{i+1}"
            val = st_session_state.get(val_key, 0)
            val_text = texto_valoracion.get(val, "")
            row.cells[2].text = f"{val}. {val_text}" if val else ""
            # Celda de condición vacía en criterios
            row.cells[3].text = ""

        # Última fila: valoración de la condición
        row = table.rows[5]
        row.cells[0].text = ""
        row.cells[1].text = ""
        row.cells[2].text = ""
        cond_key = subdim.replace(".", "_")  # Ejemplo: D1.1 → D1_1
        cond_val = st_session_state.get(cond_key, 0)
        cond_text = texto_valoracion_cond.get(cond_val, "")
        row.cells[3].text = f"{cond_val}. {cond_text}"
        set_cell_background(row.cells[3], color_puntaje.get(cond_val, "FFFFFF"))

        doc.add_paragraph("")

##############################################

def guardar_respuesta_actual():
    if "historico_respuestas" not in st.session_state:
        st.session_state.historico_respuestas = []

    # Crea una copia de las respuestas actuales y agrega un timestamp único
    copia_respuestas = st.session_state.get("respuestas", {}).copy()
    copia_respuestas["timestamp"] = datetime.now().isoformat()

    # Añade la respuesta al histórico
    st.session_state.historico_respuestas.append(copia_respuestas)


def exportar_formulario_completo_con_tablas():
    doc = Document()

    doc.add_heading('EVALUAR – BPS', level=1)
    doc.add_paragraph('EVALUACIÓN DE CONDICIONES ESENCIALES DEL ENFOQUE BIOPSICOSOCIAL EN SERVICIOS DE REHABILITACIÓN')

    # I. INFORMACIÓN DE LA INSTITUCIÓN
    doc.add_heading('I. INFORMACIÓN DE LA INSTITUCIÓN', level=2)
    campos = [
        ("Fecha", "fecha"),
        ("Departamento", "departamento"),
        ("Municipio", "municipio"),
        ("Nombre de la IPS", "nombre_institucion"),
        ("NIT", "nit"),
        ("Naturaleza jurídica", "naturaleza_juridica"),
        ("Empresa Social del Estado", "empresa_social_estado"),
        ("Nivel de atención del prestador", "nivel_atencion_prestador")
    ]
    for label, key in campos:
        valor = st.session_state.get(key, ("", ""))
        doc.add_paragraph(f"{label}: {valor[0] if isinstance(valor, tuple) else valor}")###ajuste aquí


# II. SERVICIOS DE REHABILITACIÓN HABILITADOS EN TABLA
    doc.add_heading('II. SERVICIOS DE REHABILITACIÓN HABILITADOS', level=2)

# Crear la tabla con 2 filas de encabezado: títulos de bloque + subcampos
    bloques = {
        "Servicio": [""],
        "Días de atención": ["L", "M", "Mi", "J", "V", "S", "D"],
        "Áreas de atención": ["CE", "HO", "UR", "U", "UCI", "Otr"],
        "Modalidad": ["AMB", "HOS", "DOM", "JORN", "UNMOV", "TMIA", "TMNIA", "TE", "TMO"],
        "Tipo de prestador": ["PREM", "PREF"]
    }

# Aplanar todos los encabezados finales
    headers = [h for grupo in bloques.values() for h in grupo]
    n_cols = len(headers)

# Crear tabla con 2 filas de encabezado
    table = doc.add_table(rows=2, cols=n_cols)
    table.style = 'Table Grid'

# Primera fila: títulos de bloque
    col_idx = 0
    for titulo, subcampos in bloques.items():
        colspan = len(subcampos)
        cell = table.rows[0].cells[col_idx]
        cell.text = titulo
        if colspan > 1:
            for i in range(1, colspan):
                table.rows[0].cells[col_idx + i].merge(table.rows[0].cells[col_idx])
        col_idx += colspan

# Segunda fila: subcampos
    for i, h in enumerate(headers):
        table.rows[1].cells[i].text = h

# Filas de datos (máximo 7 servicios)
    for i in range(1, 8):
        servicio = st.session_state.get(f"servicio_{i}")
        if servicio and servicio != "Seleccione":
            row = table.add_row().cells
            col = 0

        # Servicio
            row[col].text = servicio
            col += 1

        # Días
            for d in bloques["Días de atención"]:
                row[col].text = "X" if st.session_state.get(f"{d}_{i}") else ""
                col += 1

        # Áreas
            for a in bloques["Áreas de atención"]:
                row[col].text = "X" if st.session_state.get(f"{a}_{i}") else ""
                col += 1

        # Modalidad
            for m in bloques["Modalidad"]:
                row[col].text = "X" if st.session_state.get(f"{m}_{i}") else ""
                col += 1
    
        # Tipo de prestador
            row[col].text = "X" if st.session_state.get(f"prestador_P_REM_{i}", False) else ""
            col += 1
            row[col].text = "X" if st.session_state.get(f"prestador_P_REF_{i}", False) else ""
            col += 1
                


    # III. RECURSO HUMANO EN TABLA
    doc.add_heading("III. RECURSO HUMANO DE LOS SERVICIOS DE REHABILITACIÓN", level=2)
    rh_table = doc.add_table(rows=1, cols=2)
    rh_table.style = "Table Grid"
    rh_table.rows[0].cells[0].text = "Profesional"
    rh_table.rows[0].cells[1].text = "Cantidad"

    for i in range(1, 9):
        prof = st.session_state.get(f"DesP_{i}")
        cantidad = st.session_state.get(f"numero_{i}")
        if prof and prof != "Seleccione":
            row = rh_table.add_row().cells
            row[0].text = prof
            row[1].text = str(cantidad or "")

    # Aclaraciones
    aclaraciones = st.session_state.get("aclaraciones", "")
    if aclaraciones:
        doc.add_paragraph("Aclaraciones sobre la oferta de servicios o recurso humano:")
        doc.add_paragraph(aclaraciones)

    # Representantes
    doc.add_heading("Representantes de la Institución", level=2)
    for i in range(1, 7):
        rep = st.session_state.get(f"rep_inst_{i}")
        if rep:
            doc.add_paragraph(f"{i}. {rep}")

    # Profesionales verificadores
    doc.add_heading("Responsables de verificación", level=2)
    for i in range(1, 3):
        ver = st.session_state.get(f"prof_verif_{i}")
        if ver:
            doc.add_paragraph(f"{i}. {ver}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
#######################################

def exportar_primera_pagina():
    doc = Document()

    doc.add_heading('EVALUAR – BPS', level=1)
    doc.add_paragraph('EVALUACIÓN DE CONDICIONES ESENCIALES DEL ENFOQUE BIOPSICOSOCIAL EN SERVICIOS DE REHABILITACIÓN')

    # I. INFORMACIÓN DE LA INSTITUCIÓN
    doc.add_heading('I. INFORMACIÓN DE LA INSTITUCIÓN', level=2)
    for campo in ["fecha", "departamento", "municipio", "nombre_institucion", "nit", "naturaleza_juridica", "empresa_social_estado", "nivel_atencion_prestador"]:
        valor = st.session_state.get(campo, "")
        # Extrae solo la etiqueta si es una tupla
        if isinstance(valor, tuple):
            valor = valor[0]
        doc.add_paragraph(f"{campo.replace('_', ' ').capitalize()}: {valor}")

    

    # II. SERVICIOS DE REHABILITACIÓN HABILITADOS
    doc.add_heading('II. SERVICIOS DE REHABILITACIÓN HABILITADOS', level=2)
    for i in range(1, 8):
        servicio = st.session_state.get(f"servicio_{i}")
        if servicio and servicio != "Seleccione":
            doc.add_heading(f"Servicio {i}: {servicio}", level=3)

            dias = [d for d in ["L", "M", "Mi", "J", "V", "S", "D"] if st.session_state.get(f"{d}_{i}")]
            doc.add_paragraph("Días de atención: " + ", ".join(dias))

            areas = [a for a in ["CE", "HO", "UR", "U", "UCI", "Otr"] if st.session_state.get(f"{a}_{i}")]
            doc.add_paragraph("Áreas asistenciales: " + ", ".join(areas))

            modalidades = [m for m in ["AMB", "HOS", "DOM", "JORN", "UNMOV", "TMIA", "TMNIA", "TE", "TMO"] if st.session_state.get(f"{m}_{i}")]
            doc.add_paragraph("Modalidades de prestación: " + ", ".join(modalidades))

            prestador = [
                nombre for clave, nombre in [
                (f"prestador_P_REM_{i}", "P.REM"),
                (f"prestador_P_REF_{i}", "P.REF"),
                ] if st.session_state.get(clave, False)
            ]
            doc.add_paragraph("Tipo de prestador: " + ", ".join(prestador))

    # III. RECURSO HUMANO
    doc.add_heading("III. RECURSO HUMANO", level=2)
    for i in range(1, 9):
        profesional = st.session_state.get(f"DesP_{i}")
        cantidad = st.session_state.get(f"numero_{i}")
        if profesional and profesional != "Seleccione":
            doc.add_paragraph(f"{profesional}: {cantidad} profesionales")

    aclaraciones = st.session_state.get("aclaraciones", "")
    if aclaraciones:
        doc.add_paragraph("Aclaraciones sobre oferta o recurso humano:")
        doc.add_paragraph(aclaraciones)

    # REPRESENTANTES DE LA INSTITUCIÓN
    doc.add_heading("Representantes de la Institución", level=2)
    for i in range(1, 7):
        rep = st.session_state.get(f"rep_inst_{i}")
        if rep:
            doc.add_paragraph(f"{i}. {rep}")

    # PROFESIONALES RESPONSABLES DE VERIFICACIÓN
    doc.add_heading("Responsables de verificación", level=2)
    for i in range(1, 3):
        ver = st.session_state.get(f"prof_verif_{i}")
        if ver:
            doc.add_paragraph(f"{i}. {ver}")

    # Exportar buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
############################################################
#########################################################################


####   Personalización para tabla de exportación

def set_cell_background(cell, rgb_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), rgb_color)
    tcPr.append(shd)

# Diccionario de colores para cada puntaje:
color_puntaje = {
    5: '92D050',  # Verde fuerte
    4: 'C6E0B4',  # Verde medio
    3: 'FFEB9C',  # Amarillo claro
    2: 'FCE4D6',  # Naranja pálido
    1: 'F8CBAD',  # Rojo claro
}


#######################################################

doc=Document()

nombres_dimensiones = {
    "D1": "1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN",
    "D2": "2. PROCESO DE REHABILITACIÓN",
    "D3": "3. RESULTADOS DEL PROCESO DE REHABILITACIÓN"
}

nombres_subdimensiones = {
    "D1.1": "D1.1 La oferta de servicios de rehabilitación corresponde con el nivel de complejidad de la institución. ►",
    "D1.2": "D1.2 El talento humano de rehabilitación vinculado a la institución es acorde a la capacidad instalada versus la demanda de los servicios.►",
    "D1.3": "D1.3 La prestación de los servicios de rehabilitación se realiza en diferentes modalidades: intramural, extramural y/o telemedicina.",
    "D1.4": "D1.4 La institución cuenta con un sistema unificado de historia clínica disponible para los profesionales que intervienen en el proceso de rehabilitación.►",
    "D1.5": "D1.5 La atención de los usuarios de rehabilitación o “proceso de rehabilitación” se encuentra documentado en la institución.►",
    "D1.6": "D1.6 El proceso de rehabilitación se estructura por etapas o fases que orientan la atención del usuario en la institución.►",
    "D1.7": "D1.7 En los servicios de rehabilitación se encuentran disponibles guías de práctica clínica, protocolos de atención y/o procedimientos para orientar la toma de decisiones.►",
    "D1.8": "D1.8 La institución estructura e implementa un plan de capacitación en atención o rehabilitación con enfoque biopsicosocial.",
    "D1.9": "D1.9 La institución cuenta con áreas de atención, dotación y tecnología para la implementación de intervenciones orientadas a optimizar el proceso de rehabilitación.",
    "D2.1": "D2.1 Se realiza o se cuenta con valoración médica integral de la condición de salud de los usuarios de rehabilitación.",
    "D2.2": "D2.2 Se usan pruebas estandarizadas y/o instrumentos para la evaluación de los usuarios de rehabilitación.►",
    "D2.3": "D2.3 En la evaluación se valora el estado funcional del usuario.►",
    "D2.4": "D2.4 La evaluación considera el desempeño y los roles del usuario en diferentes entornos.",
    "D2.5": "D2.5 En la evaluación se identifican facilitadores y barreras del entorno que influyen en el proceso de rehabilitación del usuario.",
    "D2.6": "D2.6 En la evaluación se registran las expectativas del usuario, la familia o cuidador respecto al proceso de rehabilitación.►",
    "D2.7": "D2.7 El plan de atención del usuario de rehabilitación se estructura de acuerdo al modelo de atención y se centra en la persona.►",
    "D2.8": "D2.8 El plan de atención integra el manejo médico de la condición de salud y las intervenciones para el logro de los objetivos y/o metas de rehabilitación.",
    "D2.9": "D2.9 Los profesionales definen con el usuario, la familia y/o cuidador, objetivos y/o metas de rehabilitación que se orientan a optimizar el funcionamiento.►",
    "D2.10": "D2.10 Se establecen objetivos y/o metas de rehabilitación medibles y alcanzables en un tiempo determinado.►",
    "D2.11": "D2.11 La intervención en rehabilitación del usuario se orienta a mejorar su autonomía e independencia.►",
    "D2.12": "D2.12 Durante la intervención del usuario los profesionales de rehabilitación realizan acciones conjuntas, coordinadas e interdependientes.",
    "D2.13": "D2.13 En el proceso de rehabilitación se implementan acciones con enfoque diferencial.",
    "D2.14": "D2.14 Durante el proceso de atención, se realizan acciones para involucrar activamente al usuario, su familia y/o cuidador en el cumplimiento de los objetivos de rehabilitación.",
    "D2.15": "D2.15 En la etapa o fase de intervención se realiza reevaluación del usuario para identificar los logros y de ser necesario, realizar ajustes al plan de atención.►",
    "D2.16": "D2.16 El proceso de rehabilitación incluye acciones planificadas de orientación y canalización del usuario y su familia a otras instituciones o sectores que pueden contribuir a su participación.",
    "D2.17": "D2.17 Se realiza evaluación final del usuario para determinar los logros, y definir el egreso o la pertinencia de continuar con el proceso de rehabilitación.►",
    "D2.18": "D2.18 Se implementan acciones específicas para la atención y el egreso de usuarios de rehabilitación de larga permanencia con pobre pronóstico funcional.",
    "D3.1": "D3.1 Se utilizan instrumentos adaptados y validados en el contexto nacional para evaluar los resultados del proceso de rehabilitación.",
    "D3.2": "D3.2 Se miden y analizan los resultados del estado funcional de los usuarios posterior al proceso de rehabilitación.",
    "D3.3": "D3.3 Se mide la satisfacción de los usuarios con la atención recibida en los servicios de rehabilitación."
}



# Lista de nombres de variables en el orden deseado
orden_columnas = [
    "fecha", "departamento", "municipio", "nombre_institucion", "nit", "nombre_responsable",
    "naturaleza_juridica", "empresa_social_estado", "nivel_atencion_prestador",
    "servicio_1", "servicio_2", "servicio_3", "servicio_4", "servicio_5", "servicio_6", "servicio_7",
    # Agrega aquí el resto de keys que quieras guardar y su orden
]

# Inicializar un diccionario para almacenar los valores de los widgets con key, en el orden deseado
def extraer_variables_con_key_ordenado():
    data = {}
    for key in orden_columnas:
        if key in st.session_state:
            data[key] = st.session_state[key]
        else:
            data[key] = None
    return data

# Crear un DataFrame vacío al inicio (puedes usarlo para almacenar varias respuestas si lo deseas)
if "df_respuestas" not in st.session_state:
    st.session_state.df_respuestas = pd.DataFrame(columns=orden_columnas)


subdimension_a_paso = {
    "D1.1": 3,    "D1.2": 4,    "D1.3": 5,    "D1.4": 6,    "D1.5": 7,    "D1.6": 8,
    "D1.7": 9,    "D1.8": 10,    "D1.9": 11,    "D2.1": 12,    "D2.2": 13,    "D2.3": 14,    
    "D2.4": 15,    "D2.5": 16,    "D2.6": 17,    "D2.7": 18,    "D2.8": 19,    "D2.9": 20,    
    "D2.10": 21,    "D2.11": 22,    "D2.12": 23,    "D2.13": 24,    "D2.14": 25,    "D2.15": 26,    
    "D2.16": 27,    "D2.17": 28,    "D2.18": 29,    "D3.1": 30,    "D3.2": 31,    "D3.3": 32
}




st.session_state.pagina = 1



##########################################

def calcular_puntaje_por_dimensiones2(dimensiones, respuestas, alcance):
    puntajes = {}
    maximos = {}
    puntaje_max = 5  # pon aquí el valor máximo de tu escala

    # Usa solo las dimensiones válidas para el alcance
    if alcance == "Básico":
        # Asegúrate de definir 'dimensiones_basico' según tus reglas
        dim_usar = dimensiones_basico
    else:
        dim_usar = dimensiones

    for dim, subs in dim_usar.items():
        puntajes[dim] = 0
        maximos[dim] = 0
        for sub in subs:
            valor = respuestas.get(sub, 0)
            if valor in ("Seleccione", "No aplica", None, ""):
                valor_num = 0
            else:
                try:
                    valor_num = int(valor)
                except (TypeError, ValueError):
                    valor_num = 0
            puntajes[dim] += valor_num
            maximos[dim] += puntaje_max
    return puntajes, maximos

#################################### 


# --- Inicializar session_state ---
if "alcance" not in st.session_state:
    st.session_state.alcance = "Seleccione"
if "pagina" not in st.session_state:
    st.session_state.pagina = "info"
if "respuestas" not in st.session_state:
    st.session_state.respuestas = {}


if 'paso' not in st.session_state:
    st.session_state.paso = 1
##### para inhabilitar

def subdimension_habilitada(subdim, alcance):
    if alcance == "Completo":
        return True
    for subdim_list in dimensiones_basico.values():
        if subdim in subdim_list:
            return True
    return False




#############--------------------------------------------------------
def calcular_puntaje_por_dimensiones_v3(respuestas, alcance):
    # Define tus diccionarios fuera de la función o pásalos como argumentos si prefieres
    dimensiones_basico = {
        "D1": ["D1_1", "D1_2", "D1_4", "D1_5", "D1_6", "D1_7"],
        "D2": ["D2_2", "D2_3", "D2_6", "D2_7", "D2_9", "D2_10", "D2_11", "D2_15", "D2_17"]
    }
    dimensiones_completo = {
        "D1": ["D1_1", "D1_2", "D1_3", "D1_4", "D1_5", "D1_6", "D1_7","D1_8","D1_9"],
        "D2": ["D2_1", "D2_2", "D2_3","D2_4", "D2_5", "D2_6", "D2_7","D2_8", "D2_9","D2_10", "D2_11", "D2_12", "D2_13", "D2_14", "D2_15", "D2_16", "D2_17", "D2_18"],
        "D3": ["D3_1", "D3_2", "D3_3"]
    }
    puntaje_max = 5

    if alcance == "Básico":
        dims = dimensiones_basico
    elif alcance == "Completo":
        dims = dimensiones_completo
    else:
        dims = dimensiones_completo

    puntajes = {}
    maximos = {}

    for dim, subdims in dims.items():
        puntajes[dim] = 0
        maximos[dim] = len(subdims) * puntaje_max
        for sub in subdims:
            valor = respuestas.get(sub, 0)
            # Si la respuesta es una tupla (Texto, valor), toma el valor
            if isinstance(valor, tuple):
                val = valor[1]
            else:
                val = valor
            # Si la respuesta es "Seleccione", "No aplica", None, '', pon 0
            if val in ("Seleccione", "No aplica", None, ""):
                val = 0
            try:
                val = int(val)
            except (TypeError, ValueError):
                val = 0
            puntajes[dim] += val

    return puntajes, maximos


########## Definiendo dimensiones

dimensiones = {
    #--------------------DIMENSIÓN 1
    "D1.1": ["pD1_1_1", "pD1_1_2", "pD1_1_3", "pD1_1_4", "D1_1", "obsD1_1"],
    "D1.2": ["pD1_2_1", "pD1_2_2", "pD1_2_3", "pD1_2_4", "D1_2", "obsD1_2"],
    "D1.3": ["pD1_3_1", "pD1_3_2", "pD1_3_3", "pD1_3_4", "D1_3", "obsD1_3"],
    "D1.4": ["pD1_4_1", "pD1_4_2", "pD1_4_3", "pD1_4_4", "D1_4", "obsD1_4"],
    "D1.5": ["pD1_5_1", "pD1_5_2", "pD1_5_3", "pD1_5_4", "D1_5", "obsD1_5"],
    "D1.6": ["pD1_6_1", "pD1_6_2", "pD1_6_3", "pD1_6_4", "D1_6", "obsD1_6"],
    "D1.7": ["pD1_7_1", "pD1_7_2", "pD1_7_3", "pD1_7_4", "D1_7", "obsD1_7"],
    "D1.8": ["pD1_8_1", "pD1_8_2", "pD1_8_3", "pD1_8_4", "D1_8", "obsD1_8"],
    "D1.9": ["pD1_9_1", "pD1_9_2", "pD1_9_3", "pD1_9_4", "D1_9", "obsD1_9"],
    #---------------------DIMENSIÓN 2
    "D2.1": ["pD2_1_1", "pD2_1_2", "pD2_1_3", "pD2_1_4", "D2_1", "obsD2_1"],
    "D2.2": ["pD2_2_1", "pD2_2_2", "pD2_2_3", "pD2_2_4", "D2_2", "obsD2_2"],
    "D2.3": ["pD2_3_1", "pD2_3_2", "pD2_3_3", "pD2_3_4", "D2_3", "obsD2_3"],
    "D2.4": ["pD2_4_1", "pD2_4_2", "pD2_4_3", "pD2_4_4", "D2_4", "obsD2_4"],
    "D2.5": ["pD2_5_1", "pD2_5_2", "pD2_5_3", "pD2_5_4", "D2_5", "obsD2_5"],
    "D2.6": ["pD2_6_1", "pD2_6_2", "pD2_6_3", "pD2_6_4", "D2_6", "obsD2_6"],
    "D2.7": ["pD2_7_1", "pD2_7_2", "pD2_7_3", "pD2_7_4", "D2_7", "obsD2_7"],
    "D2.8": ["pD2_8_1", "pD2_8_2", "pD2_8_3", "pD2_8_4", "D2_8", "obsD2_8"],
    "D2.9": ["pD2_9_1", "pD2_9_2", "pD2_9_3", "pD2_9_4", "D2_9", "obsD2_9"],
    "D2.10": ["pD2_10_1", "pD2_10_2", "pD2_10_3", "pD2_10_4", "D2_10", "obsD2_10"],
    "D2.11": ["pD2_11_1", "pD2_11_2", "pD2_11_3", "pD2_11_4", "D2_11", "obsD2_11"],
    "D2.12": ["pD2_12_1", "pD2_12_2", "pD2_12_3", "pD2_12_4", "D2_12", "obsD2_12"],
    "D2.13": ["pD2_13_1", "pD2_13_2", "pD2_13_3", "pD2_13_4", "D2_13", "obsD2_13"],
    "D2.14": ["pD2_14_1", "pD2_14_2", "pD2_14_3", "pD2_14_4", "D2_14", "obsD2_14"],
    "D2.15": ["pD2_15_1", "pD2_15_2", "pD2_15_3", "pD2_15_4", "D2_15", "obsD2_15"],
    "D2.16": ["pD2_16_1", "pD2_16_2", "pD2_16_3", "pD2_16_4", "D2_16", "obsD2_16"],
    "D2.17": ["pD2_17_1", "pD2_17_2", "pD2_17_3", "pD2_17_4", "D2_17", "obsD2_17"],
    "D2.18": ["pD2_18_1", "pD2_18_2", "pD2_18_3", "pD2_18_4", "D2_18", "obsD2_18"],
    #-----------------------DIMENSIÓN 3
    "D3.1": ["pD3_1_1", "pD3_1_2", "pD3_1_3", "pD3_1_4", "D3_1", "obsD3_1"],
    "D3.2": ["pD3_2_1", "pD3_2_2", "pD3_2_3", "pD3_2_4", "D3_2", "obsD3_2"],
    "D3.3": ["pD3_3_1", "pD3_3_2", "pD3_3_3", "pD3_3_4", "D3_3", "obsD3_3"]  
}

# Agrupar automáticamente por prefijo (D1, D2, D3)
todas_dimensiones = defaultdict(list)

for subdim in dimensiones.keys():
    match = re.match(r"(D\d+)\.", subdim)
    if match:
        dimension_general = match.group(1)
        todas_dimensiones[dimension_general].append(subdim)

# Convertir a dict normal si lo prefieres
todas_dimensiones = dict(todas_dimensiones)


# Define qué subdimensiones cuentan para Básico
dimensiones_basico = {
    "D1": ["D1.1", "D1.2", "D1.4", "D1.5", "D1.6", "D1.7"],  # Según lo que mencionas
    "D2": ["D2.2", "D2.3", "D2.6", "D2.7", "D2.9", "D2.15", "D2.17"]
    # Si D3 no aplica en básico, puedes omitirla
}


# Define qué subdimensiones cuentan para Básico
dimensiones_completo = {
    "D1": ["D1.1", "D1.2", "D1.3", "D1.4", "D1.5", "D1.6", "D1.7","D1.8","D1.9"],  # Según lo que mencionas
    "D2": ["D2.1", "D2.2", "D2.3","D2.4", "D2.5", "D2.6", "D2.7","D2.8", "D2.9","D2.10", "D2.11", "D2.12", "D2.13", "D2.14", "D2.15", "D2.16", "D2.17", "D2.18"],
    "D3": ["D1.3", "D3.2", "D3.3"]
    # Si D3 no aplica en básico, puedes omitirla
}




# Ejemplo de uso: para obtener los datos actuales en un DataFrame
# df_actual = pd.DataFrame([extraer_variables_con_key_ordenado()], columns=orden_columnas)

st.markdown("""
    <style>
        .css-18e3th9 { padding: 0rem 4rem 2rem 4rem; }  /* más margen lateral interno */
        .block-container { max-width: 85%; padding-left: 2rem; padding-right: 2rem; }
        label { font-weight: 600; font-size: 1.1rem; }
    </style>
    """, unsafe_allow_html=True)



if 'historico' not in st.session_state:
    st.session_state.historico = []
##################################################################3



# Define los pasos para cada alcance
pasos_completo = list(range(1, 34)) 

pasos_basico = [2, 3, 4, 6, 7, 8, 9, 13, 14, 17, 18, 20, 21, 22, 26, 28,33]
#alcance = st.session_state.alcance_evaluacion


def pasos_validos(alcance):
    if alcance == "Básico":
        return pasos_basico
    else:
        return pasos_completo



############################################################3


if 'respuestas' not in st.session_state:
   st.session_state.respuestas = {}

def guardar_respuesta(key, value):
    if "respuestas" not in st.session_state:
        st.session_state.respuestas = {}
    st.session_state.respuestas[key] = value


if "departamento" not in st.session_state:
    st.session_state.departamento = ""

if "municipio" not in st.session_state:
    st.session_state.municipio = ""

# Si no se ha inicializado el nombre de la institución, establecer un valor por defecto
if "nombre_institucion" not in st.session_state:
    st.session_state.nombre_institucion = ""
# Si no se ha inicializado el NIT, establecer un valor por defecto
if "nit" not in st.session_state:
    st.session_state.nit = ""
# Si no se ha inicializado el nombre del responsable, establecer un valor por defecto
if "nombre_responsable" not in st.session_state:
    st.session_state.nombre_responsable = ""
# Si no se ha inicializado la naturaleza jurídica, establecer un valor por defecto
if "naturaleza_juridica" not in st.session_state:
    st.session_state.naturaleza_juridica = "Seleccione una opción..."
# Si no se ha inicializado la empresa social de estado, establecer un valor por defecto
if "empresa_social_estado" not in st.session_state:
    st.session_state.empresa_social_estado = "Seleccione una opción..."
# Si no se ha inicializado el nivel de atención del prestador, establecer un valor por defecto
if "nivel_atencion_prestador" not in st.session_state:
    st.session_state.nivel_atencion_prestador = "Seleccione una opción..."
# Si no se ha inicializado el servicio 1, establecer un valor por defecto
if "servicio_1" not in st.session_state:
    st.session_state.servicio_1 = "Seleccione"
# Si no se ha inicializado el servicio 2, establecer un valor por defecto
if "servicio_2" not in st.session_state:
    st.session_state.servicio_2 = "Seleccione"

def siguiente():
    st.session_state.paso += 1

                
def anterior():
    st.session_state.paso -= 1


opciones = [
    ("Seleccione", 0),
    ("1 - No cumple", 1),
    ("2 - Incipiente", 2),
    ("3 - Aceptable", 3),
    ("4 - Satisfactorio", 4),
    ("5 - Óptimo", 5)
]


opciones2 = [
    ("Seleccione", 0),
    ("1. No cumple no implementada", 1),
    ("2. La condición cumple de forma incipiente uno o dos críterios", 2),
    ("3. Cumple de forma aceptable mínimo tres criterios", 3),
    ("4. Cumple de forma satisfactoria mínimo tres criterios", 4),
    ("5. Cumple de forma óptima todos los críterios", 5)
]

unique_id = str(uuid.uuid4()) 
guardar_respuesta("unique_id", unique_id)  # Guarda el ID único en el estado de la sesión


if "uuid_respuesta" not in st.session_state:
    st.session_state.uuid_respuesta = str(uuid.uuid4())


st.session_state.respuestas["uuid"] = st.session_state.uuid_respuesta

####################### título y encabezado #######################
st.markdown("""
    <style>
        .block-container { padding-top: 0.1rem !important; }
    </style>
""", unsafe_allow_html=True)

st.image("Logo_ideal.png")

#st.title("EVALUAR – BPS \n  **EVALUACIÓN DE CONDICIONES ESENCIALES DEL ENFOQUE BIOPSICOSOCIAL EN SERVICIOS DE REHABILITACIÓN**")
st.markdown("""
<div style="
    background-color: #FFE066; 
    padding: 1px 8px;
    border-radius: 10px; 
    text-align: center;
    font-weight: bold;
    font-size: 1.2rem;
    line-height: 1.6;
    border: 1px solid #f0c040;
">
    EVALUAR – BPS<br>
    <span style="font-size: 1rem; padding: 1px 3px;">
        EVALUACIÓN DE CONDICIONES ESENCIALES DEL ENFOQUE BIOPSICOSOCIAL EN SERVICIOS DE REHABILITACIÓN
    </span>
</div>
""", unsafe_allow_html=True)

#--------------------------------------------------

if st.session_state.paso == 1:
#Información de la institución
    st.markdown("""
                <div style="
                background-color: #0b3c70;
                color: white;
                padding: 2px 6px;
                border-radius: 3px;
                font-size: 12px;
                font-weight: bold;
                ">
                I. INFORMACIÓN DE LA INSTITUCIÓN
                </div>
                """, unsafe_allow_html=True)


    col1, col2, col3 = st.columns([5, 1, 2])
    with col1:
        st.markdown("""
                <div style="
                background-color: #e8f0fe;
                color: black;
                padding: 2px 6px;
                font-weight: bold;
                border-radius: 0.5px;
                ">
                Diligenciar previo a la visita y validar posteriormente con los delegados de la institución.
                </div>
                """, unsafe_allow_html=True)
    
        #st.markdown("Diligencias previo a la visita y validar posteriormente con los delegados de la institución.")
    with col2:
    # Alineación vertical + espaciado elegante
        st.markdown('<div style="padding-top: 0.6rem; text-align:right;"><strong>FECHA</strong></div>', unsafe_allow_html=True)
    with col3:
    # Selector de fecha sin etiqueta visible
        fecha=st.date_input("", date.today(), label_visibility="collapsed", key="fecha")
        guardar_respuesta("fecha", fecha)
    
    col1, col2 = st.columns([4,4])
    with col1:
        st.markdown("**DEPARTAMENTO**")
        departamento=st.text_input(
            "DEPARTAMENTO", 
            value=st.session_state.respuestas.get("departamento", ""),
            label_visibility="collapsed", 
            key="departamento"
        )
        guardar_respuesta("departamento", departamento)
    with col2:
        st.markdown("**MUNICIPIO**")
        municipio=st.text_input(
            "MUNICIPIO", 
            value=st.session_state.respuestas.get("municipio",""),
            label_visibility="collapsed", 
            key="municipio"
        )
        guardar_respuesta("municipio", municipio)

    
    col1,col2 = st.columns([4, 2])
    with col1:
        st.markdown("**INSTITUCIÓN PRESTADORA DE SERVIVIOS DE SALUD**")
        nombre_institucion=st.text_input(
            "INSTITUCIÓN", 
            value=st.session_state.respuestas.get("nombre_institucion", ""),
            placeholder="Digite nombre completo del prestador", 
            label_visibility="collapsed",
            key="nombre_institucion")
        guardar_respuesta("nombre_institucion", nombre_institucion)
    with col2:
        st.markdown("**NIT**")
        nit = st.text_input(
            "NIT",
            value=st.session_state.respuestas.get("nit", ""),
            placeholder="Digite número-DV", 
            label_visibility="collapsed",
            key="nombre_responsable")
        guardar_respuesta("nit", nit)
    col1, col2, col3 = st.columns([3, 3, 3])
    with col1:
        st.markdown("**NATURALEZA JURÍDICA**")
        opcionesNJ = [("Seleccione una opción...", 0), ("Pública", 1), ("Privada", 2), ("Mixta", 3)]

        # Obtener el valor guardado o el valor por defecto (0)
        valor_guardado = st.session_state.respuestas.get("naturaleza_juridica", 0)

        # Buscar el índice según el valor guardado
        index = next((i for i, op in enumerate(opcionesNJ) if op[1] == valor_guardado), 0)

        # Mostrar el selectbox, posicionando en la opción correcta
        seleccion = st.selectbox(
            "",
            opcionesNJ,
            format_func=lambda x: x[0],
            index=index,
            key="naturaleza_juridica"
    )
        guardar_respuesta("naturaleza_juridica", seleccion[1])
        
    with col2:
        st.markdown("**EMPRESA SOCIAL DE ESTADO**")
        opciones_ese = [("Seleccione una opción...",0),("Si",1),("No",2)]
        valor_guardado = st.session_state.respuestas.get("empresa_social_estado", 0)
        index = next((i for i, op in enumerate(opciones_ese) if op[1] == valor_guardado), 0)
        seleccion = st.selectbox(
            "",
            opciones_ese,
            format_func=lambda x: x[0],
            index=index,
            key="empresa_social_estado"
        )
        guardar_respuesta("empresa_social_estado", seleccion[1])

    with col3:
        st.markdown("**NIVEL DE ATENCIÓN DEL PRESTADOR**")
        opciones_nivel = [("Seleccione una opción...",0),("1",1),("2",2),("3",3),("No aplica",4)]
        valor_guardado = st.session_state.respuestas.get("nivel_atencion_prestador", 0)
        index = next((i for i, op in enumerate(opciones_nivel) if op[1] == valor_guardado), 0)
        seleccion = st.selectbox(
            "",
            opciones_nivel,
            format_func=lambda x: x[0],
            index=index,
            key="nivel_atencion_prestador"
        )
        guardar_respuesta("nivel_atencion_prestador", seleccion[1]) 



#Información de la institución
    st.markdown("""
                <div style="
                background-color: #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 12px;
                font-weight: bold;
                ">
                II. OFERTA DE SERVICIOS DE REHABILITACIÓN
                </div>
                """, unsafe_allow_html=True)    
    
    st.markdown("""
                <div style="
                background-color: #e8f0fe ;
                color: black;
                padding: 2px 8px;
                font-weight: bold;
                border-radius: 0.5px;
                ">
                Diligenciar con los delegados de la institución.
                </div>
                
                <div style="padding: 8px; border: 1px solid #ccc; font-size: 14px;">
                <p><strong>DÍAS DE ATENCIÓN</strong> &nbsp; L: lunes &nbsp; M: martes &nbsp; Mi: miércoles &nbsp; J: jueves &nbsp; V: viernes &nbsp; S: sábado &nbsp; D: domingo</p><p><strong>ÁREA DE ATENCIÓN</strong> &nbsp; CE: Consulta externa &nbsp; HOS: Hospitalización &nbsp; UR: Urgencias &nbsp; UCI: Unidad de Cuidado Intensivo &nbsp; Otr: Otra área</p>
                <p><strong>MODALIDADES DE PRESTACIÓN</strong> &nbsp; AMB: Ambulatoria &nbsp; HOSP: Hospitalaria &nbsp; DOM: Domiciliaria &nbsp; JORN: Jornada de Salud &nbsp; UN.MOV: Unidad Móvil &nbsp; TM-IA: Telemedicina interactiva &nbsp; TM-NIA: Telemedicina no interactiva</p>
                <p><strong>TE:</strong> Teleexperticia &nbsp; <strong>TMO:</strong> Telemonitoreo</p>
                <p><strong>PRESTADOR DE TELEMEDICINA</strong> &nbsp; P.REM: Prestador remisior &nbsp; P.REF: Prestador de referencia</p>
                </div>
                """, unsafe_allow_html=True)

########### esta es la parte de los servicios la acabo de modificar.
################## HASTA AQUÍ    

    ############# CICLO BUCLE


    st.markdown(f"""
        <div style="
        background-color: #e8f0fe ;
        color: black;
        padding: 4px 10px;
        font-weight: normal;
        border-radius: 0.5px;
        "><b> SERVICIOS DE REHABILITACIÓN HABILITADOS 
        </div>
        """, unsafe_allow_html=True)
    for i in range(1, 8):
################
        if "respuestas" not in st.session_state:
            st.session_state.respuestas = {}
        opciones_servicios = [
            "Seleccione", "Fisioterapia", "Fonoaudiología", "Terapia ocupacional",
            "Terapia Respiratoria", "Esp. medicina Física y Rehabilitación",
            "Psicología", "Trabajo Social", "Nutrición","Otros profesionales"]
        
        valor_guardado = st.session_state.respuestas.get(f"servicio_{i}", "Seleccione")
        index = opciones_servicios.index(valor_guardado) if valor_guardado in opciones_servicios else 0

        servicio = st.selectbox(
            "",
            opciones_servicios,
            index=index,
            key=f"servicio_{i}"
        )

        guardar_respuesta(f"servicio_{i}", servicio)
        col_dias,sep1,col_areas, sep2,col_modalidades,col_prestador = st.columns([1,0.1,1.3,0.1,1.8,1])

###################
        
        with col_dias:
            st.markdown("<div style='text-align: center;'><b>Días de atención</b></div>", unsafe_allow_html=True)
            st.markdown("Marque con una X los días de atención")
            col1, col2, col3, col4, col5, col6, col7 = st.columns([1,1,1.2,1,1,1,1])
            dias = ["L", "M", "Mi", "J", "V", "S", "D"]
            cols = [col1, col2, col3, col4, col5, col6, col7]
            for col, dia in zip(cols, dias):
                with col:
                    st.markdown(f"**{dia}**")
                    valor = st.checkbox(
                        "",  # sin texto largo
                        value=st.session_state.respuestas.get(f"{dia}_{i}", False),
                        key=f"{dia}_{i}"
                    )
                    guardar_respuesta(f"{dia}_{i}", valor)
        with sep1:
            st.markdown("<div class='vertical-divider'></div>", unsafe_allow_html=True)
        with col_areas:
            st.markdown("<div style='text-align: center;'><b>Áreas asistenciales</b></div>", unsafe_allow_html=True)
            st.markdown("Marque con X las áreas donde se prestan servicios de rehabilitación")
            col1, col2, col3, col4, col5, col6 = st.columns(6)
            with col1:
                st.markdown("**CE**")
                area_CE = st.checkbox(
                    "",
                    value = st.session_state.respuestas.get(f"CE_{i}",False),
                    key=f"CE_{i}")
                guardar_respuesta(f"CE_{i}", area_CE)
            with col2:
                st.markdown("**HO**")
                area_HO = st.checkbox(
                    "",
                    value=st.session_state.respuestas.get(f"HO_{i}", False),
                    key=f"HO_{i}"
                )
                guardar_respuesta(f"HO_{i}", area_HO)
            with col3:
                st.markdown("**UR**")
                area_UR = st.checkbox(
                    "",
                    value=st.session_state.respuestas.get(f"UR_{i}", False),
                    key=f"UR_{i}"
                )
                guardar_respuesta(f"UR_{i}", area_UR)
            with col4:
                st.markdown("**U**")
                area_U = st.checkbox(
                    "",
                    value=st.session_state.respuestas.get(f"U_{i}", False),
                    key=f"U_{i}"
                )
                guardar_respuesta(f"U_{i}", area_U)
            with col5:
                st.markdown("**UCI**")
                area_UCI = st.checkbox(
                    "",
                    value=st.session_state.respuestas.get(f"UCI_{i}", False),
                    key=f"UCI_{i}"
                )
                guardar_respuesta(f"UCI_{i}", area_UCI)
            with col6:
                st.markdown("**Otr**")
                area_Otr = st.checkbox(
                    "",
                    value=st.session_state.respuestas.get(f"Otr_{i}", False),
                    key=f"Otr_{i}"
                )
                guardar_respuesta(f"Otr_{i}", area_Otr)
        with sep2:
            st.markdown("<div class='vertical-divider'></div>", unsafe_allow_html=True)
        with col_modalidades:
            st.markdown("<div style='text-align: center;'><b>Modalidades de prestación</b></div>", unsafe_allow_html=True)
            st.markdown("Marque con X  las modalidades habilitadas")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown("**Intramural**")
                mod_AMB = st.checkbox(
                    "AMB",
                    value=st.session_state.respuestas.get(f"AMB_{i}",False),
                    key=f"AMB_{i}")
                guardar_respuesta(f"AMB_{i}", mod_AMB)
                mod_HOS = st.checkbox(
                    "HOS",
                    value=st.session_state.respuestas.get(f"HOS_{i}",False),
                    key=f"HOS_{i}")
                guardar_respuesta(f"HOS_{i}", mod_HOS)
            with col2:
                st.markdown("**Extramural**")
                mod_DOM = st.checkbox(
                    "DOM",
                    value=st.session_state.respuestas.get(f"DOM_{i}",False),
                    key=f"DOM_{i}")
                guardar_respuesta(f"DOM_{i}", mod_DOM)
                mod_JORN = st.checkbox(
                    "JORN",
                    value=st.session_state.respuestas.get(f"JORN_{i}",False),
                    key=f"JORN_{i}")
                guardar_respuesta(f"JORN_{i}", mod_JORN)
                mod_UNMOV = st.checkbox(
                    "UN.MOV",
                    value=st.session_state.respuestas.get(f"UNMOV_{i}",False),
                    key=f"UNMOV_{i}")
                guardar_respuesta(f"UNMOV_{i}", mod_UNMOV)
            with col3:
                st.markdown("**Telemedicina**")
                mod_TMIA = st.checkbox(
                    "TM-IA",
                    value=st.session_state.respuestas.get(f"TMIA_{i}",False),
                    key=f"TMIA_{i}")
                guardar_respuesta(f"TMIA_{i}", mod_TMIA)
                mod_TMNIA = st.checkbox(
                    "TM-NIA",
                    value=st.session_state.respuestas.get(f"TMNIA_{i}",False),
                    key=f"TMNIA_{i}")
                guardar_respuesta(f"TMNIA_{i}", mod_TMNIA)
                mod_TE = st.checkbox(
                    "TE",
                    value=st.session_state.respuestas.get(f"TE_{i}",False),
                    key=f"TE_{i}")
                guardar_respuesta(f"TE_{i}", mod_TE)
                mod_TMO = st.checkbox(
                    "TMO",
                    value=st.session_state.respuestas.get(f"TMO_{i}",False),
                    key=f"TMO_{i}")
                guardar_respuesta(f"TMO_{i}", mod_TMO)
        with col_prestador:
            st.markdown("<div style='text-align: center;'><b>Prestador telemedicina</b></div>", unsafe_allow_html=True)
            st.markdown("marque con una X el tipo de prestador")
            prestador_P_REM = st.checkbox(
                "P.REM",
                value=st.session_state.respuestas.get(f"prestador_P_REM_{i}",False),
                key=f"prestador_P_REM_{i}")
            guardar_respuesta(f"prestador_P_REM_{i}", prestador_P_REM)
            prestador_P_REF = st.checkbox(
                "P.REF", 
                value=st.session_state.respuestas.get(f"prestador_P_REF_{i}",False),
                key=f"prestador_P_REF_{i}")
            guardar_respuesta(f"prestador_P_REF_{i}", prestador_P_REF)

################################ Información recursos humanos
    
    #Información de la institución
    st.markdown("""
                <div style="
                background-color: #0b3c70;
                color: white;
                padding: 2px 5px;
                border-radius: 3px;
                font-size: 12px;
                font-weight: bold;
                ">
                III. RECURSO HUMANO DE LOS SERVICIOS DE REHABILITACIÓN
                </div>
                """, unsafe_allow_html=True)    
    st.markdown("""
                <div style="
                background-color: #e8f0fe ;
                color: black;
                padding: 4px 10px;
                font-weight: normal;
                border-radius: 0.5px;
                margin-bottom: -18px;
                ">
                Registre <b>número de profesionales de los servicios de rehabilitación</b> contratado por la institución en el momento de la verificación. 
                </div>
                """, unsafe_allow_html=True)

    st.markdown("""
    <style>
    .stSelectbox div[data-baseweb="select"] {
        min-height: 30px;
    }
    input[type="number"], input[type="text"] {
        height: 30px !important;
        font-size: 14px;
    }
    /* Reduce el margen superior e inferior del selectbox */
    .stSelectbox {
        margin-top: -9px !important;
        margin-bottom: -9px !important;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <style>
    .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
    }
    .stTextInput, .stSelectbox, .stNumberInput, .stRadio {
        margin-top: -9 !important;
        margin-bottom: -9 !important;
    }
    </style>
    """, unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns([1, 1, 1, 1])




# Inicializa respuestas en session_state si no existe
    if "respuestas" not in st.session_state:
        st.session_state.respuestas = {}

    opciones = [
        "Seleccione",
        "Fisioterapia",
        "Fonoaudiología",
        "Terapia ocupacional",
        "Terapia Respiratoria",
        "Esp. medicina Física y Fehabilitación",
        "Psicología",
        "Trabajo Social",
        "Nutrición",
        "Otros profesionales",
    ]

# Para layout en 4 columnas
    cols = st.columns(4)
    pares = 8  # Número de pares selectbox/number_input

    for i in range(pares):
        select_key = f"DesP_{i+1}"
        number_key = f"numero_{i+1}"

        col = cols[i % 4]  # Distribuye en columnas

        with col:
        # Recupera y valida valor guardado para selectbox
            valor_guardado = st.session_state.respuestas.get(select_key, "Seleccione")
            if valor_guardado not in opciones:
                valor_guardado = "Seleccione"
            val = st.selectbox(
                "",
                options=opciones,
                index=opciones.index(valor_guardado),
                key=select_key,
            )
            st.session_state.respuestas[select_key] = val

        # Recupera valor guardado para number_input
            num_valor_guardado = st.session_state.respuestas.get(number_key, 0)
            num = st.number_input(
                "",
                min_value=0,
                max_value=100,
                value=num_valor_guardado,
                step=1,
                key=number_key,
            )
            st.session_state.respuestas[number_key] = num


    
    st.markdown("""
    <style>
    .titulo-caja {
        background-color: #cce5f5;
        padding: 8px;
        font-weight: bold;
        border-radius: 5px;
        font-size: 14px;
    }
    .linea {
        margin-top: 8px;
        margin-bottom: 8px;
        border: none;
        border-top: 1px solid #ddd;
    }
    </style>
    """, unsafe_allow_html=True)
    
    #st.markdown("<hr class='linea'>", unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color: #e8f0fe ;
                color: black;
                padding: 4px 10px;
                font-weight: normal;
                border-radius: 0.5px;
                ">
                Registre <b>aclaraciones pertinentes sobre la oferta de servicios de rehabilitación y el talento humano relacionado:</b> variaciones en la disponibilidad de los servicios, otras áreas donde se prestan servicios de rehabilitación. 
                </div>
                """, unsafe_allow_html=True)
    

    aclaraciones = st.text_area("",
                                value = st.session_state.respuestas.get("aclaraciones", ""),
                                height=80, 
                                key="aclaraciones")
    guardar_respuesta("aclaraciones", aclaraciones)



    st.markdown("""
                <div style="
                background-color: #0b3c70;
                color: white;
                padding: 2px 5px;
                border-radius: 3px;
                font-size: 12px;
                font-weight: bold;
                ">
                III. RECURSO HUMANO DE LOS SERVICIOS DE REHABILITACIÓN
                </div>
                """, unsafe_allow_html=True)    
    


    st.markdown("""
    <div style="
        background-color: #e8f0fe;
        border: 0px solid #ccc;
        padding: 0px 0px;
        margin-bottom: 0px;
        font-weight: bold;
        font-size: 14px;
    ">
        <h0 style='margin: 0; font-weight: bold;'>NOMBRE DE REPRESENTANTES DE LA INSTITUCIÓN [CARGO]</h0>
    </div>
    """, unsafe_allow_html=True)


    # Aplica un estilo CSS para reducir el margen inferior de los inputs
    st.markdown("""
    <style>
    .stTextInput {
        margin-top: -10px !important;
        margin-bottom: -7px !important;
        padding-bottom: 0 !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    for i in range(1, 7):
        rep = st.text_input(
            label="",
            placeholder=f"{i}. Digite nombre completo [Cargo]",
            key=f"rep_inst_{i}",
            value = st.session_state.respuestas.get(f"rep_inst_{i}", ""),
        )
        st.session_state.respuestas[f"rep_inst_{i}"] = rep
        guardar_respuesta(f"rep_inst_{i}", rep)



# 🔹 Profesionales responsables de verificación
    st.markdown("""
    <div style="
        background-color: #e8f0fe;
        border: 0px solid #ccc;
        padding: 0px 0px;
        margin-bottom: 1px;
        font-weight: bold;
        font-size: 14px;
    ">
        <h0 style='margin: 0; font-weight: bold;'>NOMBRE DE PROFESIONALES RESPONSABLES DE VERIFICACIÓN</h0>
    </div>
    """, unsafe_allow_html=True)

    #for i in range(1, 3):
    #    key = f"prof_verif_{i}"
    #    prof = st.text_input(
    #        "",
    #        placeholder=f"{i}. Digite nombre completo",
    #        key=key,
    #        value=st.session_state.respuestas.get(key, "")
    #    )
    #    st.session_state.respuestas[key] = prof
    #    guardar_respuesta(f"prof_verif_{i}", prof)

    
    for i in range(1, 4):
        prof = st.text_input(
            label= "",
            placeholder=f"{i}. Digite nombre completo", 
            key=f"prof_verif_{i}",
            value = st.session_state.respuestas.get(f"prof_verif_{i}", ""),
        )
        st.session_state.respuestas[f"prof_verif_{i}"] = prof
        guardar_respuesta(f"prof_verif_{i}", prof)


    #### botones página 1
    col1, col2, col3, col4 = st.columns([1,2,2,1])

    with col1:
        st.button("◀️ Anterior", on_click=anterior)
    with col2:
        if st.button("📄 Descargar formulario con tablas (Word)"):
            word_file = exportar_formulario_completo_con_tablas()
            st.download_button(
                label="📥 Descargar Word",
                data=word_file,
                file_name="formulario_bps_tablas.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    with col3:
        if st.button("Descargar primera página (Word)"):
            word_file = exportar_primera_pagina()
            st.download_button(
                label="📥 Descargar primera página",
                data=word_file,
                file_name="primera_pagina_formulario.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    with col4:
        st.button("Siguiente ▶️", on_click=siguiente)





##################### FORMULARIO DE EVALUACIÓN #####################
############ PÁGINA 8 #####################
elif st.session_state.paso == 2: # Evaluación de la institución.

    st.markdown("""
    <div style="background-color:#FFD966; padding: 2px 8px; font-weight:bold; border: 0px solid #b7b7b7;">
        <h0>IV. EVALUAR-BPS<h0/>
    </div>

    <div style="background-color:#DEEAF6; padding: 6px 10px; font-style:italic; border: 0px solid #b7b7b7;">
        <p style="margin: 0px;">Los siguientes ítems describen condiciones esenciales de la atención con enfoque biopsicosocial en los servicios de rehabilitación.</em></p>
        <p style="margin: 0px;">Para cada ítem los representantes de la institución deben concertar y seleccionar una respuesta entre las opciones que presenta la <strong>ESCALA DE VALORACIÓN</strong>.</em></p>
        <p style="margin: 0px;">Cada condición se acompaña de cuatro criterios de verificación para orientar la valoración.</em></p>
        <p style="margin: 0px;">Algunas condiciones serán verificadas en fuentes de información disponibles, previa autorización formal de la institución.</em></p>
    </div>

    <div style="border: 0.5px solid #b7b7b7; padding: 2 px 8px;">
        <strong>ESCALA DE VALORACIÓN</strong>
        <ul style="list-style-type: none; padding-left: 0;margin-left:8px;">
            <p style="margin: 0px;">5.</span> La condición cumple de forma óptima todos los criterios <span style="color:green; font-weight:bold;">▮</span></li>
            <p style="margin: 0px;">4.</span> La condición cumple de forma satisfactoria mínimo tres criterios</li>
            <p style="margin: 0px;">3.</span> La condición cumple de forma aceptable mínimo tres criterios</li>
            <p style="margin: 0px;">2.</span> La condición cumple de forma incipiente uno o dos criterios</li>
            <p style="margin: 0px;">1.</span> La condición no cumple ningún criterio o no se implementa <span style="color:red; font-weight:bold;">▮</span></li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

    # Al inicio del flujo, solo si no está definido
    if "alcance" not in st.session_state or st.session_state.alcance == "Seleccione":
        st.session_state.alcance = "Básico" # o el valor por defecto que prefieras


    #####################################################
    if 'alcance_seleccionado' not in st.session_state:
        st.session_state.alcance_seleccionado = False
        
    if not st.session_state.alcance_seleccionado:
        col1, col2 = st.columns([5, 1])
        with col1:
            alcance = st.radio(
                "Alcance de la evaluación:",
                options=["Básico", "Completo"],
                horizontal=True
            )
        with col2:
            if st.button("Confirmar alcance"):
                st.session_state.alcance_evaluacion = alcance
                st.session_state.alcance_seleccionado = True
                st.rerun()
    # IMPORTANTE: Aquí puedes poner un 'return' para que el usuario no vea nada más
            st.stop()
    else:
        st.markdown(f"**Alcance seleccionado:** {st.session_state.alcance_evaluacion}")
        col1, col2 = st.columns([5,1])
        with col2:
            if st.button("✏️ Modificar alcance"):
                st.session_state.alcance_seleccionado = False
                st.rerun()

            alcance = st.session_state.alcance_evaluacion
        #alcance = st.session_state.alcance_evaluacion






    

    #### botones página 2
    col1, col2= st.columns([5, 1])

    with col1:
        st.button("◀️ Anterior", on_click=anterior)
    with col2:
        st.button("Siguiente ▶️", on_click=siguiente)



elif st.session_state.paso == 3:
# Encabezado principal
    #st.markdown("### D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN")

# Descripción de la sección
# Paso 1 - D1.1
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)


    st.markdown("""
                <div style="
                background-color: #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D1.1 La oferta de servicios de rehabilitación corresponde con el nivel de complejidad de la institución.►
                </div>
                """, unsafe_allow_html=True)    
    
    with st.container():
    #st.markdown("**D1.1 La oferta de servicios de rehabilitación corresponde con el nivel de complejidad de la institución.**")
        preguntas_d11 = [
            "La institución presta servicio de psicología y/o trabajo social.",
            "La institución presta servicios de fisioterapia, fonoaudiología y/o terapia ocupacional.",
            "Los servicios de rehabilitación disponibles corresponden con el nivel de complejidad.\*",
            "Los servicios de rehabilitación se organizan en un área específica de la institución.",
        ]

        notas_d11 = [
        """Servicios de rehabilitación según nivel de atención del prestador\*:

        Nivel 3. Servicios de nivel II. Los servicios de rehabilitación se organizan en un área [Ej., unidad, departamento]. 
        Nivel 2. Medicina general y especialidades. Servicio de medicina física y rehabilitación [interconsulta], fisioterapia, 
             terapia ocupacional y/o fonoaudiología + psicología. Otras terapias y especialidades.
        Nivel 1. Medicina general o remisión de prestador externo. Servicios de fisioterapia, fonoaudiología y/o terapia ocupacional, 
             + psicología y/o trabajo social
        """]

        if notas_d11[0]:
            with st.expander("Nota"):
                st.markdown(notas_d11[0])
            

        for i, texto in enumerate(preguntas_d11):
            col1, col2= st.columns([4, 1])
            with col1:
                st.markdown(texto)
            #st.markdown("------------------------------")
    
            with col2:
                key = f"pD1_1_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D1.1:**")
            key = "D1_1"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD1_1", ""), key="obsD1_1")
            guardar_respuesta("obsD1_1", obs)
             
### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD1_1_{i+1}" for i in range(4)] + ["D1_1", "obsD1_1"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
        st.button("◀️ Anterior", on_click=anterior)
    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()




#-------------------------------------------------------------------------------------
# Paso 2 - D1.2
elif st.session_state.paso == 4:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color: #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D1.2 El talento humano de rehabilitación vinculado a la institución es acorde a la capacidad instalada versus la demanda de los servicios.►
                </div>
                """, unsafe_allow_html=True)    
    
    notas_d12 = [
    """Verificar:

    - Oportunidad de cita o atención.
    - Usuarios atendidos / hora.
    """]

    if notas_d12[0]:
        with st.expander("Nota"):
            st.markdown(notas_d12[0])

    preguntas_d12 = [
        "Los servicios de rehabilitación habilitados cuentan continuamente con profesionales contratados o vinculados. ",
        "La disponibilidad del talento humano de rehabilitación es adecuada a la capacidad instalada versus la demanda de los servicios.",
        "La institución define el perfil del talento humano de rehabilitación según las necesidades de atención de los usuarios.",
        "La institución designa un líder, coordinador o jefe de los servicios de rehabilitación.",
    ]
    for i, texto in enumerate(preguntas_d12):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("------------------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD1_2_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D1.2:**")
            key = "D1_2"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD1_2", ""), key="obsD1_2")
            guardar_respuesta("obsD1_2", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD1_2_{i+1}" for i in range(4)] + ["D1_2", "obsD1_2"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()




#-------------------------------------------------------------------------------------
# Paso 3 - D1.3
elif st.session_state.paso == 5:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color: #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D1.3 La prestación de los servicios de rehabilitación se realiza en diferentes modalidades: intramural, extramural y/o telemedicina.
                </div>
                """, unsafe_allow_html=True)    
    
    notas_d13 = [
    """ 
    """]

    if notas_d13[0]:
        with st.expander("Nota"):
            st.markdown(notas_d13[0])

    preguntas_d13 = [
        "Se prestan servicios de rehabilitación en modalidad ambulatoria y/o hospitalaria [si aplica esta modalidad].",
        "Se prestan servicios de rehabilitación en modalidad domiciliaria [u otras modalidades extramurales], y están definidos los criterios para la atención en esta[s] modalidad[es].",
        "Se prestan servicios de rehabilitación en la modalidad de telemedicina.",
        "La oferta de servicios en la modalidad de telemedicina incluye una o más especialidades médicas relacionadas con rehabilitación.",
    ]
    for i, texto in enumerate(preguntas_d13):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("------------------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD1_3_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D1.3:**")
            key = "D1_3"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD1_3", ""), key="obsD1_3")
            guardar_respuesta("obsD1_3", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD1_3_{i+1}" for i in range(4)] + ["D1_3", "obsD1_3"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



#-------------------------------------------------------------------------------------
# Paso 4 - D1.4
elif st.session_state.paso == 6:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color: #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D1.4 La institución cuenta con un sistema unificado de historia clínica disponible para los profesionales que intervienen en el proceso de rehabilitación.►
                </div>
                """, unsafe_allow_html=True)    
    
    notas_d14 = [
        """ Verificar:      

        - Historia clínica.
        - Facilitadores y barreras en la práctica.
    """]


    if notas_d14[0]:
        with st.expander("Nota"):
            st.markdown(notas_d14[0])

    preguntas_d14 = [
        "La institución cuenta con historia clínica electrónica que incluye la información del usuario en las diferentes fases de la atención.", 
        "La historia clínica incluye la atención y procedimientos de los usuarios de rehabilitación, y esta información esta disponible para los profesionales.",
        "La historia clínica está disponible en los servicios de rehabilitación para el registro simultaneo o inmediato de la atención.",
        "La historia clínica incluye contenido y/o formatos específicos para los servicios de rehabilitación.",
        ]
    for i, texto in enumerate(preguntas_d14):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("------------------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD1_4_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D1.4:**")
            key = "D1_4"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD1_4", ""), key="obsD1_4")
            guardar_respuesta("obsD1_4", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD1_4_{i+1}" for i in range(4)] + ["D1_4", "obsD1_4"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()




################ Paso 5 - D1.5
if st.session_state.paso == 7:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color: #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D1.5  La atención de los usuarios de rehabilitación o “proceso de rehabilitación” se encuentra documentado en la institución.►
                </div>
                """, unsafe_allow_html=True)    
    
    notas_d15 = [
        """ Verificar: 
        
        Documentos disponibles en Sistema de Gestión de Calidad 
        """]

    if notas_d15[0]:
        with st.expander("Nota"):
            st.markdown(notas_d15[0])

    preguntas_d15 = [
        "Se documentan los servicios de terapias y se describen: modalidades de prestación, actividades, talento humano, infraestructura, dotación, riesgos e indicadores.",
        "Se documenta la atención por rehabilitación como un proceso continuo con un tiempo de duración definido. ",
        "La documentación del proceso de rehabilitación describe los diferentes servicios que intervienen desde la entrada hasta el egreso del usuario. ",
        "El documento del proceso de rehabilitación se encuentra actualizado y disponible en el sistema de gestión de calidad.",
    ]
    for i, texto in enumerate(preguntas_d15):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("------------------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD1_5_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D1.5:**")
            key = "D1_5"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD1_5", ""), key="obsD1_5")
            guardar_respuesta("obsD1_5", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD1_5_{i+1}" for i in range(4)] + ["D1_5", "obsD1_5"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



################## Paso 6 - D1.6
elif st.session_state.paso == 8:

    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;   
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D1.6 El proceso de rehabilitación se estructura por etapas o fases que orientan la atención del usuario en la institución.►
                </div>
                """, unsafe_allow_html=True)
    notas_d16 = [
        """ Verificar:
        
        Documentos disponibles. 
        Registros de socialización.
    """]
    if notas_d16[0]:
        with st.expander("Nota"):
            st.markdown(notas_d16[0])
    preguntas_d16 = [
        "En el proceso de rehabilitación se describen los mecanismos de entrada o ingreso del usuario. ",
        "El proceso de rehabilitación se estructura por etapas o fases que orientan la atención:  1. Evaluación inicial;  2. Plan de atención; 3. Intervención y 4. Evaluación final. ",
        "En cada etapa o fase se describe el alcance y las acciones a realizar para el logro de objetivos o metas de rehabilitación.",
        "El proceso de rehabilitación se divulga al personal asistencial de la institución.",
    ]
    for i, texto in enumerate(preguntas_d16):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("------------------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD1_6_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D1.6:**")
            key = "D1_6"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD1_6", ""), key="obsD1_6")
            guardar_respuesta("obsD1_6", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD1_6_{i+1}" for i in range(4)] + ["D1_6", "obsD1_6"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



################## Paso 7 - D1.7
elif st.session_state.paso == 9:

    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;    
                font-weight: bold;
                ">
                D1.7 En los servicios de rehabilitación se encuentran disponibles guías de práctica clínica, protocolos de atención y/o procedimientos para orientar la toma de decisiones.►
                </div>
                """, unsafe_allow_html=True)
    notas_d17 = [""" Verificar:
                 
                 Protocolos de atención y GPC disponibles.
                 Procedimiento para la elaboración de GPC y protocolos.
                 Registros de socialización de GPC y protocolos.
    """]
    if notas_d17[0]:
        with st.expander("Nota"):
            st.markdown(notas_d17[0])
    preguntas_d17 = [
        "En los servicios de rehabilitación se encuentran disponibles los protocolos de atención.",
        "La institución cuenta con una o más guías de práctica clínica (GPC) específicas para rehabilitación o GPC que integran recomendaciones para rehabilitación.",
        "La institución cuenta con un procedimiento que establece la metodología para la elaboración de protocolos y GPC [metodologías: adopción, adaptación o creación].",
        "Los protocolos y/o GPC de los servicios de rehabilitación se encuentran actualizados e implementados.",
    ]
    for i, texto in enumerate(preguntas_d17):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD1_7_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D1.7:**")
            key = "D1_7"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD1_7", ""), key="obsD1_2")
            guardar_respuesta("obsD1_7", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD1_7_{i+1}" for i in range(4)] + ["D1_7", "obsD1_7"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()


################## Paso 8 - D1.8
elif st.session_state.paso == 10:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D1.8 La institución estructura e implementa un plan de capacitación en atención o rehabilitación con enfoque biopsicosocial. 
                </div>
                """, unsafe_allow_html=True)
    notas_d18 = ["""Verificar:
    
                 Registro de capacitaciones
                 Contenido de inducción y plan de capacitación
    """]
    if notas_d18[0]:
        with st.expander("Nota"):
            st.markdown(notas_d18[0])
    preguntas_d18 = [   
        "La inducción de nuevos profesionales incluye información sobre el proceso de atención con enfoque biopsicosocial.",
        "La institución realiza capacitaciones periódicas sobre la atención con enfoque biopsicosocial.",
        "Las capacitaciones sobre atención con enfoque biopsicosocial están dirigidas al personal asistencial y administrativo. [jefes, coordinadores, personal de mercadeo; RRHH].",
        "Se implementan acciones para evaluar el conocimiento del personal sobre la atención con enfoque biopsicosocial.",
    ]
    for i, texto in enumerate(preguntas_d18):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD1_8_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D1.8:**")
            key = "D1_8"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD1_8", ""), key="obsD1_8")
            guardar_respuesta("obsD1_8", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD1_8_{i+1}" for i in range(4)] + ["D1_8", "obsD1_8"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()




#################### Paso 9 - D1.9
elif st.session_state.paso == 11:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D1. ORGANIZACIÓN Y GESTIÓN DE LOS SERVICIOS DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)
    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D1.9 La institución cuenta con áreas de atención, dotación y tecnología para la implementación de intervenciones orientadas a optimizar el proceso de rehabilitación. 
                </div>
                """, unsafe_allow_html=True)
    notas_d19 = ["""Verificar:
    
                 Identificar facilitadores y barreras en la práctica [personal asistencial]. 
                 Recorrido o video.
    """]
    if notas_d19[0]:
        with st.expander("Nota"):
            st.markdown(notas_d19[0])
    preguntas_d19 = [
        "Los servicios de rehabilitación cuentan con equipos e insumos adecuados a las necesidades de la población atendida y su condición de salud.",
        "La institución realiza mantenimiento periódico y reparación oportuna de áreas, equipos e insumos de rehabilitación.",
        "Los servicios de rehabilitación disponen de tecnología que favorecen el acceso, la eficiencia y/o personalización de la atención.",
        "La institución cuenta con ambientes especializados para favorecer la autonomía, independencia y el desempeño de roles.",
    ]
    for i, texto in enumerate(preguntas_d19):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
    ### Nuevo ajuste
            with col2:
                key = f"pD1_9_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D1.9:**")
            key = "D1_9"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD1_9", ""), key="obsD1_9")
            guardar_respuesta("obsD1_9", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD1_9_{i+1}" for i in range(4)] + ["D1_9", "obsD1_9"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()


#################### Paso 10 - D2.1
elif st.session_state.paso == 12:



    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.1 Se realiza o se cuenta con valoración médica integral de la condición de salud de los usuarios de rehabilitación. 
                </div>
                """, unsafe_allow_html=True)
    notas_d21 = ["""Verificar:
    
                 Historia clínica: valoración médica
    """]
    if notas_d21[0]:
        with st.expander("Nota"):
            st.markdown(notas_d21[0])
    preguntas_d2_1 = [
        "La valoración médica de los usuarios de rehabilitación se encuentra disponible en la historia clínica.",
        "La valoración médica del usuario aborda integralmente la condición de salud para establecer el diagnóstico [diagnóstico principal y dianósticos relacionados]",
        "La información de la valoración médica es pertinente y suficiente para definir los objetivos y el plan de atención por rehabilitación.",
        "La institución cuenta con un formato estandarizado para la valoración médica de los usuarios de rehabilitación.",
    ]
    for i, texto in enumerate(preguntas_d2_1):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
    ### Nuevo ajuste
            with col2:
                key = f"pD2_1_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.1:**")
            key = "D2_1"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_1", ""), key="obsD2_1")
            guardar_respuesta("obsD2_1", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_1_{i+1}" for i in range(4)] + ["D2_1", "obsD2_1"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()

#################### Paso 11 - D2.2
elif st.session_state.paso == 13:

    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.2 Se usan pruebas estandarizadas y/o instrumentos para la evaluación de los usuarios de rehabilitación. ►
                </div>
                """, unsafe_allow_html=True)
    notas_d22 = ["""Verificar:
                 
                 Instrumento[s] de evaluación 
                 Historia clínica
    """]
    if notas_d22[0]:
        with st.expander("Nota"):
            st.markdown(notas_d22[0])
    preguntas_d2_2 = [
        "Los profesionales de rehabilitación registran en la historia clínica el uso de pruebas y/o instrumentos de evaluación.",
        "La institución define criterios para la selección y el uso de pruebas o instrumentos de evaluación de los usuarios de rehabilitación.",
        "La institución cuenta con un método desarrollado o adaptado para la evaluación de los usuarios de rehabilitación.",
        "Los profesionales hacen uso de  las pruebas o instrumentos disponibles según las caracteristicas y necesidades de los usuarios. [la disponibilidad hace referencia a fácil acceso durante la atención. Ej. en historia clínica].",
    ]
    for i, texto in enumerate(preguntas_d2_2):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_2_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.2:**")
            key = "D2_2"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_2", ""), key="obsD2_2")
            guardar_respuesta("obsD2_2", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_2_{i+1}" for i in range(4)] + ["D2_2", "obsD2_2"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()


#################### Paso 12 - D2.3
elif st.session_state.paso == 14:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.3 En la evaluación se valora el estado funcional del usuario. ►
                </div>
                """, unsafe_allow_html=True)
    notas_d23 = ["""Verificar:
                 
                 Instrumento[s] de evaluación.
                 Historia clínica.
                 **En prestadores de nivel 1: profesionales que intervienen en el proceso de rehabilitación.
    """]
    if notas_d23[0]:
        with st.expander("Nota"):
            st.markdown(notas_d23[0])
    preguntas_d2_3 = [
        "La valoración del estado funcional incluye diferentes dominios o áreas del funcionamiento de los usuarios.",
        "La valoración del estado funcional se basa en parámetros medibles y los resultados se expresan en datos numéricos y/o categóricos.",
        "La valoración del estado funcional concluye con el perfil de funcionamiento o el diagnóstico funcional del usuario.",
        "La valoración del estado funcional involucra un equipo multidisciplinario\*\* que interviene en el proceso de rehabilitación.",
    ]
    for i, texto in enumerate(preguntas_d2_3):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_3_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.3:**")
            key = "D2_3"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_3", ""), key="obsD2_3")
            guardar_respuesta("obsD2_3", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_3_{i+1}" for i in range(4)] + ["D2_3", "obsD2_3"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()

#################### Paso 13 - D2.4
elif st.session_state.paso == 15:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)
    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.4 La evaluación considera el desempeño y los roles del usuario en diferentes entornos. 
                </div>
                """, unsafe_allow_html=True)
    notas_d24 = ["""Verificar:
                 
                 Instrumento[s] de evaluación 
                 Historia clínica
                     """]
    if notas_d24[0]:
        with st.expander("Nota"):
            st.markdown(notas_d24[0])
    preguntas_d2_4 = [
        "En la evaluación se registra la ocupación o rol que desempeña el usuario en su entorno [Ej., hogar, trabajo, vida escolar].",
        "Se identifican las dificultades que presenta el usuario para el desempeño de actividades en su entorno.",
        "Se registran las expectativas del usuario y/o familia con relación a su ocupación o en el desempeño de actividades.",
        "La evaluación del usuario incluye pruebas o instrumentos para valorar la realización de actividades en su entorno.",
    ]
    for i, texto in enumerate(preguntas_d2_4):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_4_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.4:**")
            key = "D2_4"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_4", ""), key="obsD2_4")
            guardar_respuesta("obsD2_4", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_4_{i+1}" for i in range(4)] + ["D2_4", "obsD2_4"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



#################### Paso 14 - D2.5
elif st.session_state.paso == 16:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.5 En la evaluación se identifican facilitadores y barreras del entorno que influyen en el proceso de rehabilitación del usuario. 
                </div>
                """, unsafe_allow_html=True)
    notas_d25 = ["""Verificar:
                 
                 Instrumento[s] de evaluación.
                 Historia clínica.
    """]
    if notas_d25[0]:
        with st.expander("Nota"):
            st.markdown(notas_d25[0])

    st.markdown("""
                <div style="
                background-color: #f5f5f5 ;
                color: black;
                padding: 4px 10px;
                font-weight: normal;
                border-radius: 0.5px;
                "><b> En la historia clínica se registran  facilitadores y/o barreras relacionados con: 
                </div>
                """, unsafe_allow_html=True)

    preguntas_d2_5 = [
        "Acceso a servicios de salud según complejidad del diagnóstico o condición del usuario.",
        "Ayudas técnicas: disponibilidad, entrenamiento y adaptación, adecuación al entorno.",
        "Ajustes razonables en el entorno.",
        "Redes de apoyo.",
    ]
    for i, texto in enumerate(preguntas_d2_5):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("------------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_5_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.5:**")
            key = "D2_5"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_5", ""), key="obsD2_5")
            guardar_respuesta("obsD2_5", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_5_{i+1}" for i in range(4)] + ["D2_5", "obsD2_5"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



#################### Paso 15 - D2.6
elif st.session_state.paso == 17:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.6 En la evaluación se registran las expectativas del usuario, la familia o cuidador respecto al proceso de rehabilitación. ►
                </div>
                """, unsafe_allow_html=True)
    notas_d26 = ["""Verificar:
    
                 Instrumento[s] de evaluación
                 Historia clínica
                 Estrategia de acompañamiento
    """]
    if notas_d26[0]:
        with st.expander("Nota"):
            st.markdown(notas_d26[0])
    preguntas_d2_6 = [
        "La historia clínica incluye un ítem para el registro de las expectativas del usuario, la familia o cuidador.",
        "Se registran las expectativas del usuario con relación al proceso de rehabilitación.",
        "Se registran las expectativas de la familia o cuidador, especialmente en usuarios pediátricos, con compromiso cognitivo o dependencia severa.",
        "Se implementan estrategias de acompañamiento a usuarios y/o familias con expectativas no realistas frente al proceso de rehabilitación.",
    ]
    for i, texto in enumerate(preguntas_d2_6):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_6_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.6:**")
            key = "D2_6"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_6", ""), key="obsD2_6")
            guardar_respuesta("obsD2_6", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_6_{i+1}" for i in range(4)] + ["D2_6", "obsD2_6"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



#################### Paso 16 - D2.7
elif st.session_state.paso == 18:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.7 El plan de atención del usuario de rehabilitación se estructura de acuerdo al modelo de atención y se centra en la persona. ►
                </div>
                """, unsafe_allow_html=True)
    notas_d27 = ["""Verificar:
                 
                 Historia clínica
                 Plan de atención
    """]
    if notas_d27[0]:
        with st.expander("Nota"):
            st.markdown(notas_d27[0])
    preguntas_d2_7 = [
        "El plan de atención de los usuarios de rehabilitación hace parte de la historia clínica.",
        "El plan de atención tiene una estructura predeterminada que incluye los objetivos o metas de rehabilitación.",
        "En el plan de atención se describen las intervenciones a realizar por los profesionales o el equipo de rehabilitación.",
        "El plan de atención es individualizado y se basa en la condición de salud, el estado funcional, las necesidades y expectativas del usuario.",
    ]
    for i, texto in enumerate(preguntas_d2_7):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_7_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.7:**")
            key = "D2_7"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_7", ""), key="obsD2_7")
            guardar_respuesta("obsD2_7", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_7_{i+1}" for i in range(4)] + ["D2_7", "obsD2_7"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()


#################### Paso 17 - D2.8
elif st.session_state.paso == 19:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.8 El plan de atención integra el manejo médico de la condición de salud y las intervenciones para el logro de los objetivos y/o metas de rehabilitación. 
                </div>
                """, unsafe_allow_html=True)
    notas_d28 = ["""Verificar:
    
                 Historia clínica
                 Plan de atención
    """]
    if notas_d28[0]:
        with st.expander("Nota"):
            st.markdown(notas_d28[0])

    st.markdown("""
                <div style="
                background-color: #f5f5f5 ;
                color: black;
                padding: 4px 10px;
                font-weight: normal;
                border-radius: 0.5px;
                "><b> El plan de atención de los usuarios incluye: 
                </div>
                """, unsafe_allow_html=True)
    
    preguntas_d2_8 = [
        "Tratamiento médico: manejo farmacológico, procedimientos, ayudas técnicas, remisión a otros servicios [cuándo es necesario].",
        "Intervención terapéutica: terapias, psicología y otros servicios, modalidades de atención, intensidad y duración.",
        "Actividades de orientación y educación pertinentes para el usuario, la familia y/o cuidador.",
        "Actividades de canalización del usuario a servicios y/o para la gestión de apoyos que contribuyan a su participación.",
    ]
    for i, texto in enumerate(preguntas_d2_8):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_8_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.8:**")
            key = "D2_8"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_8", ""), key="obsD2_8")
            guardar_respuesta("obsD2_8", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_8_{i+1}" for i in range(4)] + ["D2_8", "obsD2_8"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



#################### Paso 18 - D2.9
elif st.session_state.paso == 20:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.9 Los profesionales definen con el usuario, la familia y/o cuidador, objetivos y/o metas de rehabilitación que se orientan a optimizar el funcionamiento. ►
                </div>
                """, unsafe_allow_html=True)
    notas_d29 = ["""Verificar:
                 
                 Historia clínica
                 Plan de atención
                 ** En prestadores de nivel 1: profesionales que intervienen en el proceso de rehabilitación.
    """]
    if notas_d29[0]:
        with st.expander("Nota"):
            st.markdown(notas_d29[0])
    preguntas_d2_9 = [
        "Los profesionales registran en la historia clínica los objetivos o metas de rehabilitación.",
        "Los objetivos y/o metas de rehabilitación están orientados a mejorar y/o potenciar la autonomía e independencia del usuario.",
        "Los profesionales involucran al usuario, la familia y/o cuidador en la definición de objetivos y/o metas de rehabilitación.",
        "Los objetivos y/o metas de rehabilitación se definen de manera concertada entre el equipo multidisciplinario,\*\* el usuario, la familia y/o cuidador.",
    ]
    for i, texto in enumerate(preguntas_d2_9):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_9_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.9:**")
            key = "D2_9"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_9", ""), key="obsD2_9")
            guardar_respuesta("obsD2_9", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_9_{i+1}" for i in range(4)] + ["D2_9", "obsD2_9"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



#################### Paso 19 - D2.10
elif st.session_state.paso == 21:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
                <div style="
                background-color:
                #0b3c70;
                color: white;
                padding: 1px 3px;
                border-radius: 3px;
                font-size: 14px;
                font-weight: bold;
                ">
                D2.10 Se establecen objetivos y/o metas de rehabilitación medibles y alcanzables en un tiempo determinado. ►
                </div>
                """, unsafe_allow_html=True)
    notas_d210 = ["""Verificar:
                  
                  Historia clínica.
                  Plan de atención.
    """]
    if notas_d210[0]:
        with st.expander("Nota"):
            st.markdown(notas_d210[0])
    preguntas_d2_10 = [
        "Los objetivos y/o metas de rehabilitación se basan en actividades funcionales alcanzables y relevantes para el usuario y/o la familia.",
        "Los objetivos y/o metas de rehabilitación son medibles y permiten determinar objetivamente los logros o resultados.",
        "En los objetivos y/o metas de rehabilitación se define un plazo o tiempo para alcanzar los logros o resultados esperados.",
        "Los objetivos y/o metas de rehabilitacion consideran la secuencialidad y progresión del proceso de rehabilitación.",
    ]
    for i, texto in enumerate(preguntas_d2_10):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_10_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.10:**")
            key = "D2_10"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_10", ""), key="obsD2_10")
            guardar_respuesta("obsD2_10", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_10_{i+1}" for i in range(4)] + ["D2_10", "obsD2_10"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()


#################### Paso 20 - D2.11
elif st.session_state.paso == 22:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    with st.container():
        st.markdown("""
                    <div style="
                    background-color:
                    #0b3c70;
                    color: white;
                    padding: 1px 3px;
                    border-radius: 3px;
                    font-size: 14px;
                    font-weight: bold;
                    ">
                    D2.11 La intervención en rehabilitación del usuario se orienta a mejorar su autonomía e independencia.  ►
                    </div>
                    """, unsafe_allow_html=True)
        notas_d211 = ["""Verificar:
                  
                      Historia clínica
                      ** En prestadores de nivel 1: profesionales que intervienen en el proceso de rehabilitación. 
        """]
        if notas_d211[0]:
            with st.expander("Nota"):
                st.markdown(notas_d211[0])

        st.markdown("""
                    <div style="
                    background-color: #f5f5f5 ;
                    color: black;
                    padding: 4px 10px;
                    font-weight: normal;
                    border-radius: 0.5px;
                    "><b> En la historia clínica de los usuarios: 
                    </div>
                    """, unsafe_allow_html=True)
    
        preguntas_d2_11 = [
            "Se registran intervenciones de rehabilitación orientadas a mejorar la realización de actividades de la vida diaria y el desempeño del usuario en su entorno.",
            "Las intervenciones de rehabilitación registradas son coherentes con los objetivos y/o metas de rehabilitación.",
            "Se registra el uso de enfoques terapéuticos, intervenciones y/o técnicas con respaldo en la evidencia.",
            "La intervención de los usuarios es realizada por el equipo multidisciplinario** e incorpora dispositivos de asistencia y tecnología.",
        ]
        for i, texto in enumerate(preguntas_d2_11):
            col1, col2 = st.columns([4, 1])
            with col1:
                st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_11_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.11:**")
            key = "D2_11"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_11", ""), key="obsD2_11")
            guardar_respuesta("obsD2_11", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_11_{i+1}" for i in range(4)] + ["D2_11", "obsD2_11"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()


        #################### Paso 21 - D2.12

elif st.session_state.paso == 23:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="
            background-color:
            #0b3c70;
            color: white;
            padding: 1px 3px;
            border-radius: 3px;
            font-size: 14px;
            font-weight: bold;
            ">
            D2.12 Durante la intervención del usuario los profesionales de rehabilitación realizan acciones conjuntas, coordinadas e interdependientes.
        </div>
        """, unsafe_allow_html=True)
    notas_d212 = ["""Verificar:
        
        Historia clínica.
        ** En prestadores de nivel 1: profesionales que intervienen en el proceso de rehabilitación. 
    """]
    if notas_d212[0]:
        with st.expander("Nota"):
            st.markdown(notas_d212[0])
    preguntas_d2_12 = [
        "Dos o más profesionales de rehabilitación de la institución intervienen al usuario de manera independiente con objetivos comunes.",
        "Los profesionales de rehabilitación realizan intervenciones disciplinares con objetivos comunes, y disponen de espacios para comunicarse y coordinar la atención.",
        "Los profesionales de rehabilitación realizan intervenciones coordinadas y complementarias con objetivos comunes, y comparten el espacio de atención.",
        "El equipo multidisciplinario\*\* dispone de espacios formales para la evaluación, seguimiento y toma de decisiones para la atención de  usuarios de mayor complejidad.",
    ]
    for i, texto in enumerate(preguntas_d2_12):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_12_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.12:**")
            key = "D2_12"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_12", ""), key="obsD2_12")
            guardar_respuesta("obsD2_12", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_12_{i+1}" for i in range(4)] + ["D2_12", "obsD2_12"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()

#################### Paso 22 - D2.13
elif st.session_state.paso == 24:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)


    st.markdown("""
        <div style="
        background-color:
        #0b3c70;
        color: white;
        padding: 1px 3px;
        border-radius: 3px;
        font-size: 14px;
        font-weight: bold;
        ">
        D2.13 En el proceso de rehabilitación se implementan acciones con enfoque diferencial. 
        </div>
        """, unsafe_allow_html=True)
    notas_d213 = ["""Verificar:
                  
                  Recorrido o video; documentación técnica; registro de capacitaciones. 
        """]
    if notas_d213[0]:
        with st.expander("Nota"):
            st.markdown(notas_d213[0])
    preguntas_d2_13 = [
        "La institución dispone de ajustes razonables para facilitar el acceso y autonomía de los usuarios con discapacidad.",
        "En la institución se cuenta con herramientas, dispositivos tecnológicos u otros mecanismos para facilitar la comunicación y participación en la toma de decisiones de los usuarios.",
        "En la institución se realizan capacitaciones al personal para brindar atención diferencial a los usuarios según su edad, género, discapacidad, etnia, orientación sexual e identidad de género.",
        "En la institución se implementan acciones diferenciadas para la atención de los usuarios según su edad, género, discapacidad, etnia, orientación sexual e identidad de género.",
    ]
    for i, texto in enumerate(preguntas_d2_13):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_13_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.13:**")
            key = "D2_13"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_13", ""), key="obsD2_13")
            guardar_respuesta("obsD2_13", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_13_{i+1}" for i in range(4)] + ["D2_13", "obsD2_13"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()




#################### Paso 23 - D2.14
elif st.session_state.paso == 25:

    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)


    st.markdown("""
        <div style="
        background-color:
        #0b3c70;
        color: white;
        padding: 1px 3px;
        border-radius: 3px;
        font-size: 14px;
        font-weight: bold;
        ">
        D2.14 Durante el proceso de atención, se realizan acciones para involucrar activamente al usuario, su familia y/o cuidador en el cumplimiento de los objetivos de rehabilitación.
        </div>
        """, unsafe_allow_html=True)
    notas_d214 = ["""Verificar:
                  
                  Historia clínica.
                  Recursos audiovisuales y contenidos.
                  Modalidades o estrategias de seguimiento o monitoreo.
        """]
    if notas_d214[0]:
        with st.expander("Nota"):
            st.markdown(notas_d214[0])
    preguntas_d2_14 = [
        "Durante la atención, los profesionales de rehabilitación brindan información al usuario y la familia sobre su rol en el proceso de rehabilitación.",
        "Los profesionales de rehabilitación entregan al usuario, la familia y/o cuidador planes de ejercicios y/o actividades para realizar en casa o en otros entornos [colegio, trabajo].",
        "En los servicios de rehabilitación se cuenta con recursos audiovisuales para informar y brindar contenido educativo a los usuarios, la familia y/o cuidador.",
        "En los servicios de rehabilitación, los profesionales disponen y hacen uso de dispositivos tecnológicos para el seguimiento o monitoreo remoto de los usuarios.",
    ]
    for i, texto in enumerate(preguntas_d2_14):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_14_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.14:**")
            key = "D2_14"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_14", ""), key="obsD2_14")
            guardar_respuesta("obsD2_14", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_14_{i+1}" for i in range(4)] + ["D2_14", "obsD2_14"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()

#################### Paso 24 - D2.15
elif st.session_state.paso == 26:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="
        background-color:
        #0b3c70;
        color: white;
        padding: 1px 3px;
        border-radius: 3px;
        font-size: 14px;
        font-weight: bold;
        ">
        D2.15 En la etapa o fase de intervención se realiza reevaluación del usuario para identificar los logros y de ser necesario, realizar ajustes al plan de atención. ►
        </div>
        """, unsafe_allow_html=True)
    notas_d215 = ["""Verificar:
    
                  Historia clínica
        """]
    if notas_d215[0]:
        with st.expander("Nota"):
            st.markdown(notas_d215[0])
    preguntas_d2_15 = [
        "Los profesionales realizan **monitoreo** continuo de signos y/o síntomas relacionados con la condición del usuario.",
        "Los profesionales registran cambios o logros en el estado funcional del paciente.",
        "Los profesionales realizan seguimiento a los objetivos de rehabilitación y hacen ajustes a la intervención cuando es necesario.",
        "La institución [o servicio] preestablece los tiempos de reevaluación de los usuarios haciendo uso de pruebas estandarizadas o instrumentos.",
    ]
    for i, texto in enumerate(preguntas_d2_15):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_15_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.15:**")
            key = "D2_15"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_15", ""), key="obsD2_15")
            guardar_respuesta("obsD2_15", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_15_{i+1}" for i in range(4)] + ["D2_15", "obsD2_15"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()


#################### Paso 25 - D2.16
elif st.session_state.paso == 27:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="
        background-color:
        #0b3c70;
        color: white;
        padding: 1px 3px;
        border-radius: 3px;
        font-size: 14px;
        font-weight: bold;
        ">
        D2.16 El proceso de rehabilitación incluye acciones planificadas de orientación y canalización del usuario y su familia a otras instituciones o sectores que pueden contribuir a su participación.
        </div>
        """, unsafe_allow_html=True)
    notas_d216 = ["""Verificar:
                  
                  -Historia clínica
                  -Documentación técnica.
        """]
    if notas_d216[0]:
        with st.expander("Nota"):
            st.markdown(notas_d216[0])
    preguntas_d2_16 = [
        "Los profesionales de rehabilitación orientan al usuario, la familia y/o cuidador sobre servicios o programas disponibles que contribuyen a la participación.",
        "Los profesionales derivan al usuario, la familia y/o cuidador a servicios o programas específicos para promover la participación del usuario. ",
        "Los servicios de rehabilitación cuentan con estrategias para la canalización del usuario y su familia a instituciones o servicios que contribuyen a la participación. ",
        "Los servicios de rehabilitación realizan trabajo en red con otras instituciones y servicios para incrementar las oportunidades de participación de los usuarios.",
    ]
    for i, texto in enumerate(preguntas_d2_16):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_16_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.16:**")
            key = "D2_16"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_16", ""), key="obsD2_16")
            guardar_respuesta("obsD2_16", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_16_{i+1}" for i in range(4)] + ["D2_16", "obsD2_16"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()

#################### Paso 26 - D2.17
elif st.session_state.paso == 28:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="
        background-color:
        #0b3c70;
        color: white;
        padding: 1px 3px;
        border-radius: 3px;
        font-size: 14px;
        font-weight: bold;
        ">
        D2.17 Se realiza evaluación final del usuario para determinar los logros, y definir el egreso o la pertinencia de continuar con el proceso de rehabilitación.►
        </div>
        """, unsafe_allow_html=True)
    notas_d217 = ["""Verificar:
                  
                  Historia clínica
        """]
    if notas_d217[0]:
        with st.expander("Nota"):
            st.markdown(notas_d217[0])
    preguntas_d2_17 = [
        "El proceso de rehabilitación de los usuarios termina con la evaluación final.",
        "Se identifican los logros o resultados según los objetivos y/o metas de rehabilitación.",
        "Con los resultados de la evaluación final, se define el egreso del usuario o la continuidad del proceso de rehabilitación.",
        "Se entregan indicaciones y recomendaciones al usuario como estrategias de mantenimiento, control médico y/o participación.",
    ]
    for i, texto in enumerate(preguntas_d2_17):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_17_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.17:**")
            key = "D2_17"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_17", ""), key="obsD2_17")
            guardar_respuesta("obsD2_17", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_17_{i+1}" for i in range(4)] + ["D2_17", "obsD2_17"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()

#################### Paso 27 - D2.18
elif st.session_state.paso == 29:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D2. PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="
        background-color:
        #0b3c70;
        color: white;
        padding: 1px 3px;
        border-radius: 3px;
        font-size: 14px;
        font-weight: bold;
        ">
        D2.18 Se implementan acciones específicas para la atención y el egreso de usuarios de rehabilitación de larga permanencia con pobre pronostico funcional.
        </div>
        """, unsafe_allow_html=True)
    notas_d218 = ["""Verificar:
    
                    Documentación técnica.
        """]
    if notas_d218[0]:
        with st.expander("Nota"):
            st.markdown(notas_d218[0])
    preguntas_d2_18 = [
        "En los servicios de rehabilitación se identifican los usuarios de larga permanencia.",
        "La institución cuenta con criterios definidos para la admisión y reingreso de los usuarios de larga permanencia.",
        "En los servicios de rehabilitación se implementan medidas específicas para la atención de los usuarios de larga permanencia.",
        "La institución establece acuerdos formales con las aseguradoras para la atención de los usuarios de larga permanencia.",
    ]
    for i, texto in enumerate(preguntas_d2_18):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD2_18_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D2.18:**")
            key = "D2_18"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD2_18", ""), key="obsD2_18")
            guardar_respuesta("obsD2_18", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD2_18_{i+1}" for i in range(4)] + ["D2_18", "obsD2_18"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()


#################### Paso 28 - D3.1
elif st.session_state.paso == 30:

    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D3. RESULTADOS DEL PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)


    st.markdown("""
        <div style="
        background-color:
        #0b3c70;
        color: white;
        padding: 1px 3px;
        border-radius: 3px;
        font-size: 14px;
        font-weight: bold;
        ">
        D3.1 Se utilizan instrumentos adaptados y validados en el contexto nacional para evaluar los resultados del proceso de rehabilitación.
        </div>
        """, unsafe_allow_html=True)
    notas_d31 = ["""Verificar;
    
                    Historia clínica; documentación técnica.
        """]
    if notas_d31[0]:
        with st.expander("Nota"):
            st.markdown(notas_d31[0])
    preguntas_d3_1 = [
        "Los instrumentos de evaluación de los usuarios de rehabilitación se encuentran validados. [priorizar instrumentos de evaluación funcional o de condiciones más frecuentes]",
        "Los requisitos o condiciones de aplicación de los instrumentos [Ej., tiempo, equipos] son viables para su uso en los servicios de rehabilitación.",
        "El uso de instrumentos de evaluación cumple con las normas de licenciamiento o derechos de autor.",
        "Los profesionales de rehabilitación reciben capacitación o entrenamiento en el uso de instrumentos de evaluación.",
    ]
    for i, texto in enumerate(preguntas_d3_1):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD3_1_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D3.1:**")
            key = "D3_1"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD3_1", ""), key="obsD3_1")
            guardar_respuesta("obsD3_1", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD3_1_{i+1}" for i in range(4)] + ["D3_1", "obsD3_1"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



#################### Paso 29 - D3.2
elif st.session_state.paso == 31:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D3. RESULTADOS DEL PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="
        background-color:
        #0b3c70;
        color: white;
        padding: 1px 3px;
        border-radius: 3px;
        font-size: 14px;
        font-weight: bold;
        ">
        D3.2 Se miden y analizan los resultados del estado funcional de los usuarios posterior al proceso de rehabilitación.
        </div>
        """, unsafe_allow_html=True)
    notas_d32 = ["""Verificar:
                 
                 Historia clínica; documentación técnica; indicadores.
        """]
    if notas_d32[0]:
        with st.expander("Nota"):
            st.markdown(notas_d32[0])
    preguntas_d3_2 = [
        "El estado funcional de los usuarios se evalúa al inicio y al final del proceso de rehabilitación.",
        "En la evaluación inicial y final del estado funcional de los usuarios se usa un método o instrumento validado.",
        "Los resultados de la evaluación inicial y final del estado funcional de los usuarios se consolidan y se analizan por la institución.",
        "La institución define indicadores de resultado relacionados con el estado funcional de los usuarios de rehabilitación.",
    ]
    for i, texto in enumerate(preguntas_d3_2):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD3_2_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D3.2:**")
            key = "D3_2"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD3_2", ""), key="obsD3_2")
            guardar_respuesta("obsD3_2", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD3_2_{i+1}" for i in range(4)] + ["D3_2", "obsD3_2"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()


elif st.session_state.paso == 32:
    st.markdown("""
                <div style="
                background-color: #F1F3F5;
                padding: 2px 6px;
            border-radius: 10px;
            font-size: 12px;
            font-weight: bold;
            color: #212529;
            margin-bottom: 0rem;
    ">
        D3. RESULTADOS DEL PROCESO DE REHABILITACIÓN
    </div>
    """, unsafe_allow_html=True)
    st.markdown("""
        <div style="
        background-color:
        #0b3c70;
        color: white;
        padding: 1px 3px;
        border-radius: 3px;
        font-size: 14px;
        font-weight: bold;
        ">
        D3.3 Se mide la satisfacción de los usuarios con la atención recibida en los servicios de rehabilitación.
        </div>
        """, unsafe_allow_html=True)
    notas_d33 = ["""Verificar:
    
                 documentación técnica; formato; informe o indicadores de satisfacción. 
        """]
    if notas_d33[0]:
        with st.expander("Nota"):
            st.markdown(notas_d33[0])
    preguntas_d3_3 = [
        "Al finalizar el proceso de rehabilitación se mide la satisfacción de los usuarios.",
        "La medición de la satisfacción de los usuarios es estandarizada y los resultados se expresan en datos numéricos y/o categorías.",
        "La evaluación de la satisfacción verifica la percepción de los usuarios sobre la oportunidad, seguridad, pertinencia y resultados de la atención.",
        "Los resultados de la satisfacción de los usuarios se consolidan, analizan y los resultados dan lugar a acciones de mejora.",
    ]
    for i, texto in enumerate(preguntas_d3_3):
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(texto)
            #st.markdown("-----------------------")
    ### Nuevo ajuste
            with col2:
                key = f"pD3_3_{i+1}"
                valor_guardado = st.session_state.respuestas.get(key, 0)  # El valor guardado, por defecto 0
                index = next((j for j, op in enumerate(opciones) if op[1] == valor_guardado), 0)
                val = st.selectbox(
                    "",
                    opciones,
                    format_func=lambda x: x[0],
                    index=index,
                    key=key
                    )
                guardar_respuesta(key, val[1])


    with st.container():
        col1, col2 = st.columns([2, 4])
        with col1:
            st.markdown("**Calificación D3.3:**")
            key = "D3_3"
            valor_guardado = st.session_state.respuestas.get(key, 0)
            index = next((j for j, op in enumerate(opciones2) if op[1] == valor_guardado), 0)

            val = st.selectbox(
                "",
                opciones2,
                format_func=lambda x: x[0],
                index=index,
                key=key
                )
            guardar_respuesta(key, val[1])
        with col2:
            obs = st.text_area("Hallazgos", value=st.session_state.respuestas.get("obsD3_3", ""), key="obsD3_3")
            guardar_respuesta("obsD3_3", obs)
             
    
    ### página 3
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

# Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    preguntas_obligatorias = [f"pD3_3_{i+1}" for i in range(4)] + ["D3_3", "obsD3_3"]
    faltan = [
        key for key in preguntas_obligatorias
        if st.session_state.respuestas.get(key, None) in (None, "", "Seleccione", 0)
    ]

    if faltan:
        st.warning("Responde todas las preguntas antes de continuar.")
    # st.write(f"Faltan: {faltan}")  # Útil para depuración

    col1, col2 = st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()

    with col2:
        if indice < len(pasos) - 1:
            if st.button("Siguiente", disabled=bool(faltan)):
                st.session_state.paso = pasos[indice + 1]
                st.rerun()



#################### Final #####################
elif st.session_state.paso == 33:
#### PUNTAJES 
    alcance = st.session_state.alcance_evaluacion
    pasos = pasos_validos(alcance)
    paso_actual = st.session_state.paso

    # Obtener el índice del paso actual en la lista de pasos válidos
    indice = pasos.index(paso_actual)

    col1, col2= st.columns([5, 1])
    with col1:
    # Botón Anterior (solo si no es el primer paso válido)
        if indice > 0:
            if st.button("Anterior"):
                st.session_state.paso = pasos[indice - 1]
                st.rerun()
###########---------------------------------FINLIZACIÓN DE FORMUARIO-----------------------------------------------------############


###################-------------------------------------
    # Asumiendo que st.session_state['respuestas'] guarda {subdim: valor}
    #alcance = st.session_state.get("alcance", "Básico")
    respuestas = st.session_state.get("respuestas", {})

    puntajes, maximos = calcular_puntaje_por_dimensiones_v3(respuestas, alcance)

    st.success("¡Formulario completado! ✅")
    st.subheader("📈 Resultados por dimensión")

    for dim in puntajes:
        st.write(f"**{dim}**: {puntajes[dim]} / {maximos[dim]}")

    st.write(f"**Puntaje Total:** {sum(puntajes.values())} / {sum(maximos.values())}")

    
##########-----------------------------------Obtención del gráfico de retroalimentación-----------------------------------############
    #total_max_global = 0
    total_global = sum(puntajes.values())
    total_max_global = sum(maximos.values())
    global_pct = round((total_global / total_max_global) * 100, 1)
  
    def graficar_nivel_implementacion(valor, show=True, figsize=(4,1)):
        import matplotlib.pyplot as plt
        import io
        rangos = list(range(0, 101, 10))
        colores = ['#7B002C', '#A11A2E', '#C63A2F', '#E76A32', '#F4A822',
                   '#FADA75', '#FCECB3', '#D6EDC7', '#A6D49F', '#4C7C2D']
        etiquetas = [f"{i+1}-{i+10}" if i != 0 else "1-10" for i in rangos[:-1]]


        fig, ax = plt.subplots(figsize=figsize, dpi=100)
        for i in range(len(colores)):
            left = rangos[i]
            width = 10
            ax.barh(0, width=width, left=left, color=colores[i], edgecolor='white')
            label = f"{left+1}-{left+10}" if left != 0 else "1-10"
            ax.text(left + width/2, 0.6, label, ha='center', va='bottom', fontsize=9)
        #ax.plot(valor, 0, 'o', markersize=15, markeredgecolor='black', markerfacecolor='none')
        #ax.text(valor, 0, f'{valor:.1f}', ha='center', va='center', fontsize=7, weight='bold')

        # 1. Determina en qué rango cae el valor
        rango_idx = min(int(valor)//10, 9)  # 0 a 9
        left = rangos[rango_idx]
        right = left + 10
        centro = left + 5

    # 2. Dibuja un rectángulo debajo del número para destacarlo (opcional)
        #ax.add_patch(plt.Rectangle((left, -0.15), 10, 0.3, alpha=0.2, zorder=2))

    # 3. Escribe el número centrado en ese recuadro
        ax.text(centro, 0, f"{valor:.1f}%", ha='center', va='center', fontsize=14, color='black', weight='bold', zorder=3)
    



        ax.set_xlim(0, 100)
        ax.set_ylim(-1.2, 1.2)  # 👈 Ajusta aquí para recortar el espacio arriba
        ax.axis('off')

        img_buffer = io.BytesIO()
        fig.tight_layout(pad=0.2)  # 👈 Mejora aún más el recorte
        plt.savefig(img_buffer, format='png', bbox_inches='tight', dpi=100)
        if show:
            import streamlit as st
            st.pyplot(fig)
        plt.close(fig)
        img_buffer.seek(0)
        return img_buffer
        
    
#-------------------------------------------------------------------------------------------------------------------------------#    
#-----------------------------Llamar esta función al final con el puntaje global como porcentaje--------------------------------#

    img_buffer = graficar_nivel_implementacion(global_pct, show=True, figsize=(8,2))

################### para descargar la gráfica

    st.download_button(
        label="Descargar imagen",
        data=img_buffer,
        file_name="grafica.png",
        mime="image/png"
    )
    img_buffer.seek(0)

#----------------------------------- DEFINIR SEPARADOR PARA LOS ARCHIVOS EN EXCEL----------------------------------------------# 
    separador = st.radio(
        "Separador del archivo CSV:",
        options=[",", ";"],
        format_func=lambda x: "Coma (,)" if x == "," else "Punto y coma (;)",
        horizontal=True
    )


    # Filtrar subdimensiones que sí existen en el session_state
    resumen = []
    for sub, variables in dimensiones.items():
        if sub in nombres_subdimensiones and variables:
            #codificacion = sub
            nombre = nombres_subdimensiones[sub]
            valor_raw = st.session_state.respuestas.get(variables[4], 0)
            valor = valor_raw[1] if isinstance(valor_raw, tuple) else valor_raw 
            obs_key = variables[5] if len(variables) > 5 else None
            hallazgos = st.session_state.respuestas.get(obs_key, "") if obs_key else "No aplica"
        
            resumen.append({
                #"Código": codificacion,
                "Condición": nombre,
                "Valoración": valor,
                "Hallazgos": hallazgos
            })

    df_resumen = pd.DataFrame(resumen)
    csv_resumen = df_resumen.to_csv(index=False, sep=separador, encoding="utf-8-sig").encode("utf-8-sig")
    st.download_button(
            label="📥 Descargar resumen (CSV)",
            data=csv_resumen,
            file_name="valoracion_por_subdimension.csv",
            mime="text/csv"
            )




    # ---------------------------------------- Exportar respuestas -------------------------------------------------
#    import pandas as pd
#    # Convertir respuestas en DataFrame y exportar
#    df_respuestas = pd.DataFrame([st.session_state.respuestas])
#    csv = df_respuestas.to_csv(index=False, sep=separador, encoding="utf-8-sig").encode("utf-8-sig")

#    st.download_button(
#        label="📥 Descargar respuestas (CSV)",
#        data=csv,
#        file_name="respuestas_formulario.csv",
#        mime="text/csv"
#    )
    
    
    ruta_base = "respuestas_consolidadas.csv"

    # Convertir la respuesta actual a DataFrame
    df_actual = pd.DataFrame([st.session_state.respuestas])

    # Cargar archivo si ya existe
    if os.path.exists(ruta_base):
        df_existente = pd.read_csv(ruta_base)

        # Verificar si el UUID ya está presente
        if st.session_state.uuid_respuesta not in df_existente["uuid"].values:
            df_total = pd.concat([df_existente, df_actual], ignore_index=True)
            df_total.to_csv(ruta_base, index=False, sep=separador)
        else:
            df_total = df_existente  # Ya estaba guardado, no agregamos
    else:
        df_total = df_actual
        df_total.to_csv(ruta_base, index=False, sep=separador)

    st.download_button(
    label="📥 Descargar base acumulada (CSV)",
    data=df_total.to_csv(index=False, sep=separador, encoding="utf-8-sig").encode("utf-8-sig"),
    file_name="respuestas_consolidadas.csv",
    mime="text/csv"
    )
    df_actual = pd.DataFrame([st.session_state.respuestas])
    # Subir la respuesta actual a Google Sheets
    #subir_respuesta_a_drive(st.session_state.respuestas)



# Agrupa por dimensión
    from collections import defaultdict
    subdims_por_dim = defaultdict(list)
    for sub in dimensiones.keys():
        dim = sub.split(".")[0]  # "D1", "D2", etc.
        subdims_por_dim[dim].append(sub)

    from collections import defaultdict

# Agrupa subdimensiones por dimensión (D1, D2, D3)
    subdims_por_dim = defaultdict(list)
    for sub in dimensiones.keys():
        dim = sub.split(".")[0]
        subdims_por_dim[dim].append(sub)

    if alcance == "Básico":
        dimensiones_actuales = {
            "D1": ["D1.1", "D1.2", "D1.4", "D1.5", "D1.6", "D1.7"],
            "D2": ["D2.2", "D2.3", "D2.6", "D2.7", "D2.9", "D2.15", "D2.17"]
        }
    elif alcance == "Completo":
        dimensiones_actuales = {
            "D1": ["D1.1", "D1.2", "D1.3", "D1.4", "D1.5", "D1.6", "D1.7","D1.8","D1.9"],
            "D2": ["D2.1", "D2.2", "D2.3","D2.4", "D2.5", "D2.6", "D2.7","D2.8", "D2.9","D2.10", "D2.11", "D2.12", "D2.13", "D2.14", "D2.15", "D2.16", "D2.17", "D2.18"],
            "D3": ["D3.1", "D3.2", "D3.3"]
        }    
    else:
        dimensiones_actuales = {
            "D1": ["D1.1", "D1.2", "D1.3", "D1.4", "D1.5", "D1.6", "D1.7","D1.8","D1.9"],
            "D2": ["D2.1", "D2.2", "D2.3","D2.4", "D2.5", "D2.6", "D2.7","D2.8", "D2.9","D2.10", "D2.11", "D2.12", "D2.13", "D2.14", "D2.15", "D2.16", "D2.17", "D2.18"],
            "D3": ["D3.1", "D3.2", "D3.3"]
        }


    for dim, subdim_list in dimensiones_actuales.items():
        nombre_largo = nombres_dimensiones.get(dim, dim)
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
    

    # Fila 0: nombre largo en celda combinada y fondo gris oscuro
        titulo_row = table.rows[0]
        titulo_cell = titulo_row.cells[0]
        titulo_cell.merge(titulo_row.cells[1])
        p = titulo_cell.paragraphs[0]
        run = p.add_run(nombre_largo)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255,255,255)
        set_cell_background(titulo_cell, '4F4F4F')  # Gris oscuro

    # Fila 1: encabezados
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = 'CONDICIONES'
        hdr_cells[1].text = 'CALIFICACIÓN'
        for cell in hdr_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True


    
        for sub in subdim_list:
            mask = df_resumen["Condición"].str.contains(nombres_subdimensiones[sub], case=False, regex=False)
            if not mask.any():
                continue
            row = df_resumen[mask].iloc[0]
            val = int(row["Valoración"])
            row1 = table.add_row().cells
            row1[0].text = row["Condición"]
            row1[1].text = str(val)
            set_cell_background(row1[1], color_puntaje.get(val, 'FFFFFF'))
            row2 = table.add_row().cells
            merged = row2[0].merge(row2[1])
            merged.text = f"Hallazgos: {row['Hallazgos']}"

    # Total de la dimensión
        row_total = table.add_row().cells
        cell_dim = row_total[0]
        cell_puntaje = row_total[1]
        run_dim = cell_dim.paragraphs[0].add_run(f"TOTAL")
        run_dim.bold = True
        run_puntaje = cell_puntaje.paragraphs[0].add_run(f"{puntajes[dim]}")
        run_puntaje.bold = True

        doc.add_paragraph("")  # Salto de línea entre tablas

    
    # ... después de crear doc = Document() y antes de guardar en buffer:
    # ... continúa con el resto de tu exportación Word
    word_file_2 = tabla_detalle_condiciones(doc, dimensiones, nombres_subdimensiones, st.session_state.respuestas)
    st.download_button(
            label="📥 Descargar Word",
            data=word_file,
            file_name="formulario_bps_tablas.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# Agregar salto de página y el gráfico
    doc.add_page_break()
    doc.add_heading("📈 Nivel de Implementación Global", level=2)

# Crear gráfico

    # En la pestaña final (paso 33)

    doc.add_picture(img_buffer, width=Inches(6.5)) 

    # Guardar Word en buffer
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)

    # Botón de descarga
    st.download_button(
        label="📥 Descargar resumen (Word)",
        data=word_buffer,
        file_name="resumen_valoracion.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

#-------------------------------------------------------------------------------------------------------------------------------#
#-------------------------------------------------------------------------------------------------------------------------------#

    
    def enviar_por_correo(destinatario, asunto, cuerpo, word_buffer):
        usuario = "tata.sanchez.10@gmail.com"
        contraseña = st.secrets["correo_gmail"]
    # Guarda el archivo temporalmente
        with open("resumen_valoracion.docx", "wb") as f:
            f.write(word_buffer.getbuffer())
        yag = yagmail.SMTP(usuario, contraseña)
        yag.send(
            to=destinatario,
            subject=asunto,
            contents=cuerpo,
            attachments=["resumen_valoracion.docx"]
        )
        yag.close()
        os.remove("resumen_valoracion.docx")

# En Streamlit
    st.subheader("📧 Enviar informe por correo")
    destinatario = st.text_input("Correo destinatario")
    if st.button("Enviar informe Word", key="btn_enviar_word"):
        if destinatario:
            try:
                word_buffer.seek(0)
                enviar_por_correo(
                    destinatario,
                    "Informe del piloto",
                    "Adjunto el informe Word generado del formulario.",
                    word_buffer
                )
                st.success("¡Correo enviado con éxito!")
            except Exception as e:
                st.error(f"Ocurrió un error al enviar el correo: {e}")
        else:
            st.warning("Por favor ingresa un correo válido.")
    
###########---------------------------------------------###################

    if st.button("🏠 Volver al inicio", type="primary"):
        guardar_respuesta_actual()

        for key in list(st.session_state.keys()):
            del st.session_state[key]
    
        st.rerun()

##########---------------------------------------------#####################
############################################################################


